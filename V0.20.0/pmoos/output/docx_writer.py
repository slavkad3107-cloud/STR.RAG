"""МОДУЛЬ 5 (часть 1): формирование откорректированного раздела ПМООС в .docx.

Архитектурное замечание (важно):
  По требованию пользователя исходные файлы проекта НЕ хранятся приложением
  (см. fix #9 — храним только имя проекта и чанки/токены в RAG-базе). Поэтому
  «откорректированный ПМООС» формируется как профессиональный документ-носитель
  корректировок: для каждого принятого ответа выводится конкретная правка в
  раздел ПМООС со ссылкой на источник (раздел/файл/страница). Если пользователь
  передаёт путь к исходному файлу ПМООС (original_oos_path), его текст
  извлекается и добавляется отдельным приложением для удобства сверки —
  непосредственного слепого переписывания произвольного документа не делаем,
  чтобы не повредить нормоконтроль.

Используется python-docx (работает на машине пользователя без Node.js).
Применяются принципы оформления из docx-skill: шрифт Arial, явные ширины
колонок таблиц (DXA), нумерация средствами Word, без юникод-«буллетов».
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..paths import project_paths
from .common import (
    accepted_answers, final_answer_text, source_ref, set_default_font,
    add_title, add_heading,
)


def _add_intro(doc: Document, project: str, object_type: str, n_corr: int, accepted_only: bool) -> None:
    from .. import VERSION
    p = doc.add_paragraph()
    run = p.add_run(
        f"Документ сформирован автоматически системой {VERSION} "
        f"{datetime.now().strftime('%d.%m.%Y %H:%M')}. "
        f"Проект: «{project}». Тип объекта: {object_type}. "
    )
    run.font.size = Pt(10)
    p2 = doc.add_paragraph()
    note = (
        f"Включено корректировок: {n_corr} (только принятые пользователем)."
        if accepted_only else
        f"Включено корректировок: {n_corr} (ВНИМАНИЕ: показаны предлагаемые ответы — "
        f"ни один пункт ещё не принят пользователем в Модуле 4)."
    )
    r2 = p2.add_run(note)
    r2.font.size = Pt(10)
    r2.italic = True
    if not accepted_only:
        r2.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)


def _add_corrections(doc: Document, answers: list[dict]) -> None:
    add_heading(doc, "1. Корректировки раздела ПМООС по замечаниям экспертизы", level=1)
    if not answers:
        doc.add_paragraph("Принятых корректировок нет.")
        return
    for a in answers:
        num = a.get("number", "?")
        add_heading(doc, f"Замечание №{num}", level=2)

        pr = doc.add_paragraph()
        pr.add_run("Замечание эксперта: ").bold = True
        pr.add_run(a.get("remark", "") or "—")

        corr = (a.get("user_answer") or a.get("correction") or "").strip()
        pc = doc.add_paragraph()
        pc.add_run("Вносимая правка в ПМООС: ").bold = True
        pc.add_run(corr or "—")

        ans = final_answer_text(a)
        if ans:
            pa = doc.add_paragraph()
            pa.add_run("Ответ для экспертизы: ").bold = True
            pa.add_run(ans)

        src = source_ref(a)
        ps = doc.add_paragraph()
        rs = ps.add_run(f"Источник: {src}")
        rs.italic = True
        rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        md = (a.get("missing_data") or "").strip()
        if md:
            pm = doc.add_paragraph()
            rm = pm.add_run(f"Требуется дополнить данными: {md}")
            rm.font.size = Pt(9)
            rm.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)


def _add_original_appendix(doc: Document, original_oos_path: str | Path, cfg) -> None:
    from ..ingest.loaders import extract_file
    try:
        pages = extract_file(Path(original_oos_path), ocr=False)
    except Exception as exc:  # noqa: BLE001
        doc.add_paragraph(f"(Не удалось прочитать исходный ПМООС: {exc})")
        return
    doc.add_page_break()
    add_heading(doc, "Приложение А. Исходный текст раздела ПМООС (для сверки)", level=1)
    note = doc.add_paragraph()
    rn = note.add_run(
        "Ниже приведён извлечённый текст исходного (неоткорректированного) раздела. "
        "Используйте его как основу: примените к нему правки из раздела 1."
    )
    rn.italic = True
    rn.font.size = Pt(9)
    for pg in pages:
        txt = (pg.get("text") or "").strip()
        if not txt:
            continue
        loc = pg.get("loc", "")
        if loc:
            h = doc.add_paragraph()
            rh = h.add_run(str(loc))
            rh.bold = True
            rh.font.size = Pt(9)
            rh.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for para in txt.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)


def build_corrected_oos_docx(project: str, *, original_oos_path: str | Path | None = None,
                             cfg=None, out_path: str | Path | None = None) -> Path:
    """Сформировать .docx с откорректированным разделом ПМООС.

    Возвращает путь к созданному файлу (по умолчанию в out-папке проекта).
    """
    from ..config import load_config
    cfg = cfg or load_config()
    data = _load_answers(project)
    object_type = data.get("object_type") or cfg.get("object_type", "площадной")

    answers, accepted_only = accepted_answers(data)

    doc = Document()
    set_default_font(doc, "Arial", 11)
    add_title(doc, "ОТКОРРЕКТИРОВАННЫЙ РАЗДЕЛ ПМООС/ООС")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("Перечень мероприятий по охране окружающей среды")
    rsub.font.size = Pt(12)
    rsub.bold = True

    _add_intro(doc, project, object_type, len(answers), accepted_only)
    _add_corrections(doc, answers)

    if original_oos_path:
        _add_original_appendix(doc, original_oos_path, cfg)

    out_dir = project_paths(project)["out"]
    out_dir.mkdir(parents=True, exist_ok=True)
    out = Path(out_path) if out_path else out_dir / f"ПМООС_откорректированный_{project}.docx"
    doc.save(str(out))
    return out


def _load_answers(project: str) -> dict[str, Any]:
    from ..pipeline.block1_answers import load_answers
    return load_answers(project)


# ───────────── №10-11..14: правки ПРЯМО в исходных томах (жёлтым) ─────────────

def _anchor_token(text: str) -> str | None:
    """Маркер места из текста правки: «табл. 4.1», «т. 8.3», «раздел 5», «п. 2.3»."""
    import re as _re
    m = _re.search(r"(?:табл(?:ица|ице|ицу|\.)?|т\.|разд(?:ел|еле|\.)?|п(?:ункт|\.)\.?)"
                   r"\s*№?\s*([\d][\d.]*)", (text or "").lower())
    return m.group(1).rstrip(".") if m else None


def _insert_paragraph_after(par, runs):
    """Вставляет новый абзац СРАЗУ ПОСЛЕ par. runs = [(text, bold, yellow)]."""
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    from docx.enum.text import WD_COLOR_INDEX
    new_p = par._p.makeelement(qn("w:p"), {})
    par._p.addnext(new_p)
    np = Paragraph(new_p, par._parent)
    for t, bold, hl in runs:
        r = np.add_run(t)
        r.bold = bold
        if hl:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return np


def write_corrected_volumes(project: str, sources: list) -> list[Path]:
    """РЕАЛЬНО откорректированные тома ООС: открываем ИСХОДНЫЙ .docx, вставляем
    правки по принятым/правленым ответам с ЖЁЛТОЙ заливкой — по якорю
    («табл./п./раздел N») сразу после нужного абзаца, иначе — в конец, в раздел
    «КОРРЕКТИРОВКИ ПО ЗАМЕЧАНИЯМ ЭКСПЕРТИЗЫ». Документ НЕ пересобирается,
    поэтому большие тома (десятки МБ) обрабатываются быстро. Если томов
    несколько — ответы раскладываются по полю «Том ООС» (без тома — в первый).
    Возвращает пути файлов *_КОРР.docx."""
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    data = _load_answers(project)
    answers = [a for a in data.get("answers", [])
               if a.get("status") in ("accepted", "edited")]
    srcs = [Path(s) for s in sources if s]
    out_dir = project_paths(project)["out"]
    out_dir.mkdir(parents=True, exist_ok=True)
    outs: list[Path] = []
    if not srcs:
        return outs

    def _match(a, src: Path) -> bool:
        v = (a.get("oos_volume") or "").lower().strip()
        if not v:
            return False
        n, stem = src.name.lower(), src.stem.lower()
        if v in n or n in v or stem in v:
            return True
        vp = Path(v)
        # stem берём ТОЛЬКО если v — имя файла с настоящим расширением
        # (иначе Path("том 6.1").stem == "том 6" и правка утекает в чужой том)
        if vp.suffix.lower() in (".docx", ".doc", ".pdf"):
            return vp.stem.lower() in n
        return False

    matched_ids = {id(a) for s2 in srcs for a in answers if _match(a, s2)}
    for si, src in enumerate(srcs):
        if len(srcs) > 1:
            mine = [a for a in answers if _match(a, src)]
            if si == 0:
                mine += [a for a in answers if id(a) not in matched_ids]
        else:
            mine = list(answers)
        doc = Document(str(src))
        ptexts = [(p, (p.text or "").lower()) for p in doc.paragraphs]
        tail = []
        for a in mine:
            num = a.get("number", "?")
            ans = (a.get("user_answer") or a.get("answer") or "").strip()
            corr = (a.get("correction") or "").strip()
            runs = [(f"[Изменение по замечанию №{num}] ", True, True)]
            if ans:
                runs.append((f"ОТВЕТ: {ans} ", False, True))
            if corr:
                runs.append((f"ВНОСИМАЯ ПРАВКА: {corr}", False, True))
            tok = _anchor_token(corr) or _anchor_token(a.get("remark", ""))
            target = None
            if tok:
                for p, lt in ptexts:
                    if tok in lt:
                        target = p
                        break
            if target is not None:
                _insert_paragraph_after(target, runs)
            else:
                tail.append(runs)
        if tail:
            h = doc.add_paragraph()
            hr = h.add_run("КОРРЕКТИРОВКИ ПО ЗАМЕЧАНИЯМ ЭКСПЕРТИЗЫ")
            hr.bold = True
            for runs in tail:
                p = doc.add_paragraph()
                for t, b, hl in runs:
                    r = p.add_run(t)
                    r.bold = b
                    if hl:
                        r.font.highlight_color = WD_COLOR_INDEX.YELLOW
        out = out_dir / f"{src.stem}_КОРР.docx"
        doc.save(str(out))
        outs.append(out)
        print(f"[m5] {src.name}: правок {len(mine)} → {out.name}", flush=True)
    return outs
