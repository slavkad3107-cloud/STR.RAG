"""Извлечение текста из файлов ПД (PDF/DOCX/DOC/XLSX/XLS/TXT).

Учтены замечания ревью:
  * PDF открываем ОДИН раз (а не на каждой странице);
  * OCR запускаем, если текста на странице меньше порога (по умолчанию 200
    символов) — «многие сканы дают мусорный текст»;
  * тяжёлые библиотеки импортируем лениво.

Возвращаемая единица — «страница» документа:
  {"loc": "стр. 12" | "лист 'Расчёт'" | "абз. 40", "text": "...", "is_table": bool}
loc нужен для провенанса «Источник: файл, страница».
"""
from __future__ import annotations

import hashlib
import sqlite3
import threading
from pathlib import Path

Page = dict  # {"loc": str, "text": str, "is_table": bool}

# --- дисковый кэш OCR (оптимизация М1) --------------------------------------
# OCR (pytesseract) — самая дорогая операция ingest для сканов. Кэшируем по
# sha256(язык + байты изображения страницы): повторная индексация (после
# очистки индекса, или одинаковые сканы в разных проектах) не гоняет OCR заново.
# Файл общий между процессами (фоновый индексатор — отдельный процесс); WAL
# делает одно переиспользуемое соединение безопасным.
_OCR_LOCK = threading.Lock()
_OCR_CON: sqlite3.Connection | None = None
_OCR_MAX_ROWS = 200_000  # ограничение размера кэша OCR (FIFO-вытеснение старых записей)


def _ocr_db() -> sqlite3.Connection:
    global _OCR_CON
    if _OCR_CON is None:
        from ..paths import data_root
        con = sqlite3.connect(str(data_root() / "ocr_cache.sqlite"), check_same_thread=False)
        try:
            con.execute("PRAGMA journal_mode=WAL")
            con.execute("PRAGMA synchronous=NORMAL")
            con.execute("PRAGMA busy_timeout=5000")
        except Exception:  # noqa: BLE001
            pass
        con.execute("CREATE TABLE IF NOT EXISTS ocr (k TEXT PRIMARY KEY, v TEXT)")
        _OCR_CON = con
    return _OCR_CON


def _ocr_cache_get(key: str) -> str | None:
    try:
        with _OCR_LOCK:
            row = _ocr_db().execute("SELECT v FROM ocr WHERE k=?", (key,)).fetchone()
        return row[0] if row else None
    except Exception:  # noqa: BLE001
        return None


_OCR_PUTS = 0  # счётчик вставок: FIFO-проверку размера делаем раз в 200 put'ов


def _ocr_cache_put(key: str, val: str) -> None:
    global _OCR_PUTS
    try:
        with _OCR_LOCK:
            con = _ocr_db()
            con.execute("INSERT OR REPLACE INTO ocr (k, v) VALUES (?, ?)", (key, val))
            # FIFO-ограничение размера: COUNT(*) на КАЖДЫЙ put дорог на большом
            # кэше — проверяем периодически (переполнение между проверками ≤200 строк)
            _OCR_PUTS += 1
            if _OCR_PUTS % 200 == 0:
                (n,) = con.execute("SELECT COUNT(*) FROM ocr").fetchone()
                if n > _OCR_MAX_ROWS:
                    con.execute(
                        "DELETE FROM ocr WHERE rowid IN "
                        "(SELECT rowid FROM ocr ORDER BY rowid LIMIT ?)",
                        (n - _OCR_MAX_ROWS,))
            con.commit()
    except Exception:  # noqa: BLE001
        pass


def _ocr_page(pix_bytes: bytes, lang: str) -> str:
    key = hashlib.sha256(lang.encode("utf-8") + b"\x00" + pix_bytes).hexdigest()
    cached = _ocr_cache_get(key)
    if cached is not None:
        return cached
    try:
        import io
        from PIL import Image
        import pytesseract
        img = Image.open(io.BytesIO(pix_bytes))
        txt = pytesseract.image_to_string(img, lang=lang) or ""
    except Exception:
        return ""
    _ocr_cache_put(key, txt)
    return txt


def extract_pdf(path: Path, *, ocr: bool = True, min_text_chars: int = 200,
                lang: str = "rus+eng", max_pages: int = 0) -> list[Page]:
    pages: list[Page] = []
    try:
        import fitz  # PyMuPDF
    except Exception as e:  # pragma: no cover
        raise RuntimeError("Не установлен PyMuPDF (pip install pymupdf)") from e

    # ДВУХФАЗНАЯ обработка: фаза 1 — последовательный проход fitz (рендер pixmap
    # только здесь: PyMuPDF НЕ потокобезопасен), собираем текст и PNG страниц,
    # которым нужен OCR; фаза 2 — OCR ПАРАЛЛЕЛЬНО (tesseract — внешний процесс,
    # GIL не мешает; последовательный OCR простаивал 60-70% ядер: 300-страничный
    # скан занимал 12-27 минут). Результаты детерминированы (те же PNG-байты →
    # тот же вывод), порядок страниц сохраняется, OCR-кэш потокобезопасен (_OCR_LOCK).
    doc = fitz.open(str(path))
    total = doc.page_count
    page_text: dict[int, str] = {}
    ocr_jobs: list[tuple[int, bytes]] = []
    empty_text_pages: set[int] = set()  # чистые сканы (нет текстового слоя вовсе)
    try:
        for i, page in enumerate(doc, start=1):
            # защита от «зависания» на гигантских сканах: не рендерим больше лимита
            if max_pages and i > max_pages:
                print(f"[loaders] {path.name}: лимит {max_pages} стр. — остальные "
                      f"{total - max_pages} стр. пропущены (ocr.max_pages)", flush=True)
                break
            text = page.get_text("text") or ""
            if not text.strip():
                empty_text_pages.add(i)
            page_text[i] = text
            if ocr and len(text.strip()) < min_text_chars:
                try:
                    pix = page.get_pixmap(dpi=200)
                    ocr_jobs.append((i, pix.tobytes("png")))
                except Exception as e:  # noqa: BLE001 — раньше молча глоталось
                    print(f"[loaders] рендер стр. {i} в {path.name}: {e}", flush=True)
    finally:
        doc.close()

    if ocr_jobs:
        import os as _os
        from concurrent.futures import ThreadPoolExecutor
        workers = min(4, max(1, (_os.cpu_count() or 2) - 1), len(ocr_jobs))
        print(f"[loaders] OCR {path.name}: {len(ocr_jobs)} стр., "
              f"потоков {workers}…", flush=True)
        done = 0

        def _job(item):
            i, png = item
            try:
                return i, _ocr_page(png, lang)
            except Exception as e:  # noqa: BLE001
                print(f"[loaders] OCR стр. {i} в {path.name}: {e}", flush=True)
                return i, ""

        with ThreadPoolExecutor(max_workers=workers) as ex:
            for i, ocr_text in ex.map(_job, ocr_jobs):
                if len(ocr_text.strip()) > len(page_text.get(i, "").strip()):
                    page_text[i] = ocr_text
                done += 1
                if done % 20 == 0:
                    print(f"[loaders] OCR {path.name}: {done}/{len(ocr_jobs)}…", flush=True)

    for i in sorted(page_text):
        if page_text[i].strip():
            pages.append({"loc": f"стр. {i}", "text": page_text[i], "is_table": False})

    # таблицы — открываем pdfplumber ОДИН раз (лимит страниц тот же, что и у OCR:
    # иначе защита от гигантских PDF была бы неполной — extract_tables прошёл бы всё).
    # Страницы БЕЗ текстового слоя (чистые сканы) пропускаем: extract_tables берёт
    # текст ячеек из chars — на скане они дали бы только пустые ячейки, а парсинг
    # каждой такой страницы стоит 0.5-3 с.
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                if max_pages and i > max_pages:
                    break
                if i in empty_text_pages:
                    continue
                try:
                    for tbl in (page.extract_tables() or []):
                        rows = [" | ".join((c or "").strip() for c in row) for row in tbl]
                        ttext = "\n".join(r for r in rows if r.strip())
                        if ttext.strip():
                            pages.append({"loc": f"стр. {i} (таблица)",
                                          "text": ttext, "is_table": True})
                except Exception:
                    continue
                finally:
                    try:  # освободить page-кэш pdfplumber (рост RAM на больших PDF)
                        page.flush_cache()
                    except Exception:  # noqa: BLE001
                        pass
    except Exception:
        pass
    return pages


def extract_docx(path: Path) -> list[Page]:
    try:
        import docx  # python-docx
    except Exception as e:  # pragma: no cover
        raise RuntimeError("Не установлен python-docx") from e
    d = docx.Document(str(path))
    pages: list[Page] = []
    buf: list[str] = []
    para_idx = 0
    for p in d.paragraphs:
        para_idx += 1
        t = (p.text or "").strip()
        if t:
            buf.append(t)
        # группируем по ~1500 символов в одну «единицу»
        if sum(len(x) for x in buf) > 1500:
            pages.append({"loc": f"абз. ~{para_idx}", "text": "\n".join(buf), "is_table": False})
            buf = []
    if buf:
        pages.append({"loc": f"абз. ~{para_idx}", "text": "\n".join(buf), "is_table": False})
    # таблицы
    for ti, table in enumerate(d.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join((c.text or "").strip() for c in row.cells))
        ttext = "\n".join(r for r in rows if r.strip())
        if ttext.strip():
            pages.append({"loc": f"таблица {ti}", "text": ttext, "is_table": True})
    return pages


def extract_xlsx(path: Path) -> list[Page]:
    try:
        import openpyxl
    except Exception as e:  # pragma: no cover
        raise RuntimeError("Не установлен openpyxl") from e
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    pages: list[Page] = []
    for ws in wb.worksheets:
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            # бьём большие листы на куски
            chunk: list[str] = []
            for r in rows:
                chunk.append(r)
                if sum(len(x) for x in chunk) > 1800:
                    pages.append({"loc": f"лист '{ws.title}'", "text": "\n".join(chunk), "is_table": True})
                    chunk = []
            if chunk:
                pages.append({"loc": f"лист '{ws.title}'", "text": "\n".join(chunk), "is_table": True})
    wb.close()
    return pages


def extract_file(path: Path, *, ocr: bool = True, min_text_chars: int = 200,
                 lang: str = "rus+eng", max_pages: int = 0) -> list[Page]:
    """Диспетчер по расширению. .doc/.xls (старые бинарные) — мягко предупреждаем."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path, ocr=ocr, min_text_chars=min_text_chars, lang=lang,
                           max_pages=max_pages)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".xlsx", ".xlsm"):
        return extract_xlsx(path)
    if ext in (".txt", ".md", ".csv"):
        try:
            return [{"loc": "файл", "text": path.read_text(encoding="utf-8", errors="ignore"),
                     "is_table": ext == ".csv"}]
        except Exception:
            return []
    if ext in (".doc", ".xls"):
        # старые форматы требуют конвертации (LibreOffice/textract) — вне MVP
        raise RuntimeError(
            f"Формат {ext} не поддерживается напрямую. Сконвертируйте в "
            f"{'docx' if ext == '.doc' else 'xlsx'} (например, через MS Office/LibreOffice)."
        )
    raise RuntimeError(f"Неизвестный тип файла: {ext}")


SUPPORTED_EXT = {".pdf", ".docx", ".xlsx", ".xlsm", ".txt", ".md", ".csv"}
