"""МОДУЛЬ 5 (часть 2): таблица ответов на ВСЕ замечания.

Требование пользователя: на основе файла замечаний сформировать таблицу со
столбцами «ОТВЕТ» и «ИСТОЧНИК» (раздел/файл/страница) по каждому замечанию.
Делаем в двух форматах: .docx (для сдачи) и .xlsx (для работы).

Берём ВСЕ замечания из answers.json (по одному пункту на замечание — ничего не
теряем). Если пользователь принял/поправил ответ — выводим его, иначе выводим
предложение ИИ с пометкой статуса.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

from ..paths import project_paths
from .common import (
    final_answer_text, source_ref, source_ref_lines, set_default_font,
    add_title, shade_cell, set_repeat_header, set_col_widths, cell_text,
)

_STATUS_RU = {
    "accepted": "принято",
    "edited": "отредактировано",
    "rejected": "отклонено",
    "proposed": "предложено ИИ",
}

_HEADERS = ["№", "ТИП", "ТОМ ООС", "ЗАМЕЧАНИЕ", "ОТВЕТ", "ИСТОЧНИК", "СТАТУС"]
_WIDTHS_IN = [0.45, 1.0, 1.15, 2.9, 3.55, 1.95, 0.8]  # сумма ≈ 11.8" (альбомная Letter, поля 1")
_HEAD_FILL = "1F3B5B"


def _answers(project: str) -> dict[str, Any]:
    from ..pipeline.block1_answers import load_answers
    return load_answers(project)


# ─────────────────────────────── DOCX ───────────────────────────────
def build_answers_table_docx(project: str, *, out_path: str | Path | None = None) -> Path:
    data = _answers(project)
    rows = data.get("answers", [])

    doc = Document()
    set_default_font(doc, "Arial", 10)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.7))

    add_title(doc, "ОТВЕТЫ НА ЗАМЕЧАНИЯ ГОСЭКСПЕРТИЗЫ К РАЗДЕЛУ ПМООС")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(
        f"Проект: «{project}». Всего замечаний: {len(rows)}. "
        f"Сформировано: {datetime.now().strftime('%d.%m.%Y')}."
    )
    rm.font.size = Pt(9)

    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(_HEADERS):
        shade_cell(hdr.cells[i], _HEAD_FILL)
        cell_text(hdr.cells[i], h, bold=True, size=10, color=(255, 255, 255))

    for a in rows:
        cells = table.add_row().cells
        cell_text(cells[0], str(a.get("number", "")), size=10)
        cell_text(cells[1], a.get("category", "") or "—", size=8)
        cell_text(cells[2], a.get("oos_volume", "") or "—", size=8)
        cell_text(cells[3], a.get("remark", "") or "—", size=9)
        ans = final_answer_text(a)
        corr = (a.get("correction") or "").strip()
        body = ans if ans else "—"
        if corr and corr not in body:
            body = f"{body}\n\nПравка в ПМООС: {corr}" if ans else f"Правка в ПМООС: {corr}"
        cell_text(cells[4], body, size=9)
        cell_text(cells[5], "\n".join(source_ref_lines(a)), size=8, color=(0x44, 0x44, 0x44))
        cell_text(cells[6], _STATUS_RU.get(a.get("status", ""), a.get("status", "")), size=9)

    set_col_widths(table, _WIDTHS_IN)

    out_dir = project_paths(project)["out"]
    out_dir.mkdir(parents=True, exist_ok=True)
    out = Path(out_path) if out_path else out_dir / f"Ответы_на_замечания_{project}.docx"
    doc.save(str(out))
    return out


# ─────────────────────────────── XLSX ───────────────────────────────
def build_answers_table_xlsx(project: str, *, out_path: str | Path | None = None) -> Path:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    data = _answers(project)
    rows = data.get("answers", [])

    wb = Workbook()
    ws = wb.active
    ws.title = "Ответы на замечания"

    head_font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    head_fill = PatternFill("solid", fgColor=_HEAD_FILL)
    wrap = Alignment(wrap_text=True, vertical="top")
    thin = Side(style="thin", color="BBBBBB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for c, h in enumerate(_HEADERS, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = head_font
        cell.fill = head_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border

    r = 2
    for a in rows:
        ans = final_answer_text(a)
        corr = (a.get("correction") or "").strip()
        body = ans if ans else ""
        if corr and corr not in body:
            body = f"{body}\n\nПравка в ПМООС: {corr}" if ans else f"Правка в ПМООС: {corr}"
        values = [
            str(a.get("number", "")),
            a.get("category", "") or "",
            a.get("oos_volume", "") or "",
            a.get("remark", "") or "",
            body or "—",
            "\n".join(source_ref_lines(a)),
            _STATUS_RU.get(a.get("status", ""), a.get("status", "")),
        ]
        for c, v in enumerate(values, 1):
            cell = ws.cell(row=r, column=c, value=v)
            cell.alignment = wrap
            cell.border = border
        r += 1

    widths = [6, 15, 20, 46, 56, 32, 15]
    from openpyxl.utils import get_column_letter
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    out_dir = project_paths(project)["out"]
    out_dir.mkdir(parents=True, exist_ok=True)
    out = Path(out_path) if out_path else out_dir / f"Ответы_на_замечания_{project}.xlsx"
    wb.save(str(out))
    return out


def build_all(project: str) -> dict[str, Path]:
    """Сформировать обе таблицы. Возвращает {'docx':…, 'xlsx':…}."""
    return {
        "docx": build_answers_table_docx(project),
        "xlsx": build_answers_table_xlsx(project),
    }
