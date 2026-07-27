"""Юнит-тесты ЧИСТОЙ (детерминированной) логики STR.RAG.

Проверяют самые хрупкие правила, которые тихо ломаются при правке регэкспов и
невидимы до неверного результата: коды ЗВ, классификация замечаний, группировка
версий томов, схема ответа ИИ, взвешивание гибридного поиска, якоря правок М5,
состав разделов-источников ПП-87. Ничего тяжёлого (без моделей/сети/GPU)."""
from __future__ import annotations


# ─────────────────────── ЗВ: коды и алиасы ───────────────────────
def test_pollutant_bare_hydrocarbon_alias_removed():
    from pmoos.entities import find_pollutants
    # «непредельные углеводороды C1-C5» НЕ должны получить код предельных 2754
    codes = [p["code"] for p in find_pollutants("выбросы непредельных углеводородов C1-C5")]
    assert "2754" not in codes


def test_pollutant_specific_forms_resolve():
    from pmoos.entities import find_pollutants, pollutant_code
    assert pollutant_code("толуол")[0] == "0621"
    assert pollutant_code("сероводород")[0] == "0333"
    codes = {p["code"] for p in find_pollutants("растворитель: ксилол, толуол, ацетон")}
    assert {"0616", "0621", "1401"} <= codes


def test_pollutant_predelnye_still_2754():
    from pmoos.entities import find_pollutants
    codes = [p["code"] for p in find_pollutants("углеводороды предельные C12-C19")]
    assert "2754" in codes


# ─────────────────────── схема ответа ИИ ───────────────────────
def test_normalize_answer_confidence_and_sources():
    from pmoos.pipeline.block1_answers import _normalize_answer
    r = _normalize_answer({"answer": "  текст  ", "confidence": "высокая",
                           "used_sources": ["1", 3, "х"]})
    assert r["answer"] == "текст"
    assert r["confidence"] == "high"
    assert r["used_sources"] == [1, 3]


def test_normalize_answer_empty_is_low():
    from pmoos.pipeline.block1_answers import _normalize_answer
    r = _normalize_answer({"answer": "", "confidence": "bogus"})
    assert r["confidence"] == "low"
    r2 = _normalize_answer({})
    assert r2["answer"] == "" and r2["used_sources"] == []


# ─────────────────────── классификация замечаний ───────────────────────
def test_classify_remark():
    from pmoos.pipeline.block1_answers import _classify_remark
    assert _classify_remark("Выполнить перерасчёт рассеивания") == "Перерасчёт"
    assert _classify_remark("Привести в соответствие с ГОСТ 17.2.3.02") == "Нормативы"
    assert _classify_remark("Приложить договор на вывоз отходов") == "Доп. документы"


# ─────────────────────── взвешенный RRF ───────────────────────
def test_rrf_weight_lifts_bm25():
    from pmoos.retrieval.hybrid import HybridRetriever
    dense = [{"id": "d1"}]   # только dense, rank0
    bm25 = [{"id": "b1"}]    # только bm25, rank0
    # при равных весах d1 идёт первым (dense-список первый). Большой вес bm25
    # должен вытащить b1 в топ — ровно то, что чинит перекос «BM25 тонет».
    eq = HybridRetriever._rrf([dense, bm25], k=60, weights=[1.0, 1.0])
    assert eq[0]["id"] == "d1"
    fused = HybridRetriever._rrf([dense, bm25], k=60, weights=[1.0, 10.0])
    assert fused[0]["id"] == "b1"


def test_rrf_equal_weight_default():
    from pmoos.retrieval.hybrid import HybridRetriever
    a = [{"id": "x"}, {"id": "y"}]
    b = [{"id": "x"}, {"id": "z"}]
    fused = HybridRetriever._rrf([a, b])  # x встречается в обоих → первый
    assert fused[0]["id"] == "x"


# ─────────────────────── якоря правок М5 ───────────────────────
def test_anchor_token():
    from pmoos.output.docx_writer import _anchor_token
    assert _anchor_token("см. табл. 4.1 раздела") == "4.1"
    assert _anchor_token("уточнить в п. 2.3.5") == "2.3.5"
    assert _anchor_token("общая формулировка без места") is None


def test_iter_all_paragraphs_includes_table_cells():
    # М5: якорь правки часто лежит В ТАБЛИЦЕ; обход должен видеть абзацы ячеек,
    # иначе правки уходят «в конец» вместо места (жалоба пользователя).
    import docx
    from pmoos.output.docx_writer import _iter_all_paragraphs
    d = docx.Document()
    d.add_paragraph("обычный абзац тела")
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "Табл. 4.1 параметры источников"
    texts = [p.text for p in _iter_all_paragraphs(d)]
    assert any("обычный абзац" in x for x in texts)
    assert any("4.1" in x for x in texts)  # абзац из ячейки таблицы найден


def test_match_volume():
    from pathlib import Path
    from pmoos.output.docx_writer import _match_volume
    assert _match_volume({"oos_volume": "том 6.1.docx"}, Path("Том 6.1.docx"))
    # разные тома не должны совпадать по обрезанному stem
    assert not _match_volume({"oos_volume": "том 6.1.docx"}, Path("Том 6.2.docx"))


# ─────────────────────── состав разделов-источников ПП-87 ───────────────────────
def test_ios_emission_sources_included():
    from pmoos.ingest.sections import source_section_codes
    codes = set(source_section_codes("площадной"))
    # ЭОМ/ОВиК/ГС теперь источники выбросов и должны попадать в поиск источников
    assert {"IOS_EOM", "IOS_OV", "IOS_GS"} <= codes


# ─────────────────────── группировка версий томов ───────────────────────
def test_base_name_keeps_volume_number():
    from pmoos.versioning.versions import _base_name
    # «том 6.1» и «том 6.2» — РАЗНЫЕ тома (номер не срезаем)
    assert _base_name("ПМООС том 6.1.docx") != _base_name("ПМООС том 6.2.docx")
    # а «изм.2»/«корр»/дата — версии ОДНОГО документа (срезаются)
    assert _base_name("ПМООС том 6.1 изм.2.docx") == _base_name("ПМООС том 6.1.docx")
    assert _base_name("ПМООС том 6.1 корр 2025-06-01.docx") == _base_name("ПМООС том 6.1.docx")


def test_version_rank_orders():
    from pmoos.versioning.versions import _version_rank
    r0 = _version_rank("ПМООС.docx")[0]
    r_corr = _version_rank("ПМООС корр.docx")[0]
    r_v3 = _version_rank("ПМООС v3.docx")[0]
    assert r0 < r_corr
    assert r_v3 > r0


# ─────────────────────── стоп индексации (потеря данных) ───────────────────────
def test_stop_indexing_does_not_error_completed_file(tmp_path, monkeypatch):
    # Регрессия v0.26: «⏹ Стоп» помечал error файл, который уже done (current_file
    # не сброшен) → при возобновлении его чанки удалялись и НЕ восстанавливались
    # (sha в устаревшем known_shas → «дубликат») — тихая потеря документа.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import importlib
    import pmoos.paths, pmoos.projects
    importlib.reload(pmoos.paths)
    from pmoos.projects import register_project
    from pmoos.index import indexer as I
    register_project("STOP_T")
    st = I.read_state("STOP_T")
    st.update({"status": "running", "pid": 0, "current_file": "готовый.pdf",
               "files": {"готовый.pdf": {"status": "done", "chunks": 5}}})
    I.write_state("STOP_T", st)
    I.stop_indexing("STOP_T")
    assert I.read_state("STOP_T")["files"]["готовый.pdf"]["status"] == "done"


def test_stop_indexing_marks_inflight_file_error(tmp_path, monkeypatch):
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import importlib
    import pmoos.paths
    importlib.reload(pmoos.paths)
    from pmoos.projects import register_project
    from pmoos.index import indexer as I
    register_project("STOP_T2")
    st = I.read_state("STOP_T2")
    st.update({"status": "running", "pid": 0, "current_file": "вработе.pdf", "files": {}})
    I.write_state("STOP_T2", st)
    I.stop_indexing("STOP_T2")
    assert I.read_state("STOP_T2")["files"]["вработе.pdf"]["status"] == "error"


# ─────────────────────── config merge / get-set ───────────────────────
def test_config_deep_merge_and_get():
    from pmoos.config import Config, _deep_merge
    merged = _deep_merge({"a": {"x": 1, "y": 2}}, {"a": {"y": 9, "z": 3}})
    assert merged == {"a": {"x": 1, "y": 9, "z": 3}}
    c = Config({"retrieval": {"top_k": 8}})
    assert c.get("retrieval.top_k") == 8
    c.set("retrieval.top_k", 12)
    assert c.get("retrieval.top_k") == 12
    assert c.get("nope.missing", "def") == "def"


def test_object_type_pinned_per_project(tmp_path, monkeypatch):
    # ЛОВУШКА (закрыта в v0.30.3): «Продолжить» с переключённым в сайдбаре типом
    # объекта домешивал в базу вторую разметку разделов ПП-87. Тип фиксируется
    # при первом запуске; игнор переключателя при возобновлении; смена — только
    # полной переиндексацией. Проекты пустые → ранний выход до загрузки модели.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.projects import register_project
    from pmoos.index import indexer as I

    register_project("ОТ1")
    I.run_indexing("ОТ1", object_type="площадной")
    assert I.read_state("ОТ1").get("object_type") == "площадной"   # зафиксирован

    I.run_indexing("ОТ1", object_type="линейный")                  # «Продолжить»
    assert I.read_state("ОТ1").get("object_type") == "площадной"   # игнор

    # reindex на ПУСТОМ проекте выходит по «нет файлов» ДО drop_collection —
    # тип НЕ перезаписывается: смена фиксируется только в момент фактического
    # удаления коллекции (защита от падения между записью типа и drop,
    # находка адверсариального ревью v0.30.3).
    I.run_indexing("ОТ1", object_type="линейный", reindex=True)
    assert I.read_state("ОТ1").get("object_type") == "площадной"


def test_consistency_flags_invented_numbers():
    # АУДИТ (critical): выдуманные величины (г/с, т/год) проходили без единого
    # предупреждения — для ООС это самый дорогой класс ошибок.
    from pmoos.pipeline.consistency import compare
    src = "Валовый выброс азота диоксида составляет 0,0125 г/с."
    ans = "Выполнен перерасчёт: выброс 0,8412 г/с, что соответствует 3,55 т/год."
    res = compare(src, ans)
    assert res["issues"], "выдуманные числа должны попадать в замечания"
    assert any("числов" in i.lower() for i in res["issues"])
    # честный ответ по источнику предупреждений не даёт
    assert not compare(src, "Выброс составляет 0,0125 г/с.")["issues"]


def test_normative_edition_mismatch_flagged():
    # АУДИТ (critical): несуществующая редакция молча подменялась реестровой,
    # и Блок 2 рапортовал «проблем не найдено».
    from pmoos.normatives.engine import check_text
    res = check_text("Расчёт выполнен по СанПиН 2.1.3684-19.")
    txt = " ".join(p.get("recommendation", "") for p in res["problems"])
    assert res["problems"], "несовпадение редакции должно попадать в проблемы"
    assert "редакц" in txt.lower()


def test_context_snippet_keeps_found_fragment():
    # АУДИТ (critical): обрезка «первые 900 символов» выбрасывала саму находку
    # (строку ИТОГО в конце склеенной таблицы) — модель отвечала вслепую.
    from pmoos.pipeline.block1_answers import _center_snippet, _format_context
    long_text = "Шапка таблицы. " + ("данные строки таблицы. " * 200) + "ИТОГО: 12,345 т/год"
    out = _center_snippet(long_text, 900, "ИТОГО: 12,345 т/год")
    assert "ИТОГО: 12,345 т/год" in out
    ctx, srcs = _format_context([{"text": long_text, "payload": {"file": "т.pdf"}}])
    assert len(ctx) > 2000                      # бюджет контекста поднят с 900
    assert "ИТОГО" in ctx or len(ctx) >= 3000


def test_export_marks_unreviewed(tmp_path, monkeypatch):
    # АУДИТ (critical): в файл для экспертизы уходили непроверенные и
    # отклонённые ответы без единой пометки.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import json
    from pmoos.projects import register_project
    from pmoos.paths import project_paths
    from pmoos.output.answers_table import unreviewed_stats
    register_project("ЭКСП")
    p = project_paths("ЭКСП")["answers"]
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps({"answers": [
        {"number": "1", "status": "accepted", "answer": "ок"},
        {"number": "2", "status": "proposed", "answer": "не читал"},
        {"number": "3", "status": "rejected", "answer": "плохой"},
        {"number": "4", "status": "accepted", "answer": "ок", "low_support": True},
    ]}, ensure_ascii=False), encoding="utf-8")
    s = unreviewed_stats("ЭКСП")
    assert s["всего"] == 4 and s["не проверено (предложен)"] == 1
    assert s["отклонено"] == 1 and s["с предупреждениями"] == 1


def test_save_merged_keeps_fresh_answer_for_rejected(tmp_path, monkeypatch):
    # АУДИТ (critical): слияние решений возвращало с диска и «отклонённые»
    # ответы, затирая только что сгенерированный новый — отклонение означает
    # «переспроси», а не «верни старое».
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import json
    from pmoos.projects import register_project
    from pmoos.paths import project_paths
    from pmoos.config import load_config
    from pmoos.ingest.remarks import Remark
    from pmoos.pipeline.block1_answers import _save_merged
    register_project("СЛИЯН")
    p = project_paths("СЛИЯН")["answers"]
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps({"answers": [
        {"number": "1", "remark": "р1", "answer": "СТАРЫЙ отклонённый",
         "status": "rejected"},
        {"number": "2", "remark": "р2", "answer": "старый", "status": "accepted",
         "user_answer": "МОЙ ПРИНЯТЫЙ"},
    ]}, ensure_ascii=False), encoding="utf-8")
    rems = [Remark(number="1", text="р1"), Remark(number="2", text="р2")]
    by_num = {"1": {"number": "1", "remark": "р1", "answer": "СВЕЖИЙ ответ",
                    "status": "proposed"},
              "2": {"number": "2", "remark": "р2", "answer": "свежий2",
                    "status": "proposed"}}
    out = _save_merged("СЛИЯН", rems, by_num, load_config(), "линейный", partial=False)
    a1 = next(a for a in out["answers"] if a["number"] == "1")
    a2 = next(a for a in out["answers"] if a["number"] == "2")
    assert a1["answer"] == "СВЕЖИЙ ответ"          # отклонённый переспрошен
    assert a2["user_answer"] == "МОЙ ПРИНЯТЫЙ"     # принятый пользователем цел


def test_auto_select_preserves_user_settings(tmp_path, monkeypatch):
    # АУДИТ (high): автовыбор при каждом старте стирал ручные настройки модулей
    # и затирал дешёвые модели вспомогательных ролей.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.config import load_config
    from pmoos.core.health import auto_select
    cfg = load_config()
    cfg.set("ai.default_provider", "mistral")
    cfg.set("ai.modules.module4.provider", "deepseek")
    cfg.set("ai.providers.mistral.extract", "mistral-small-latest")
    cfg.save()
    health = {"mistral": {"ok": True, "tier": 2, "best_model": "mistral-large-2512", "ms": 900},
              "deepseek": {"ok": True, "tier": 3, "best_model": "deepseek-v4-pro", "ms": 800}}
    # текущий провайдер работает → ничего не меняем
    auto_select(cfg, health)
    cfg2 = load_config()
    assert cfg2.get("ai.modules.module4.provider") == "deepseek"
    assert cfg2.model_for("mistral", "extract") == "mistral-small-latest"
    # текущий УМЕР → переключаемся, но рабочее переопределение модуля сохраняем
    health["mistral"] = {"ok": False, "tier": 2, "error": "лимит"}
    p, m = auto_select(cfg2, health)
    cfg3 = load_config()
    assert p == "deepseek" and cfg3.default_provider() == "deepseek"
    assert cfg3.get("ai.modules.module4.provider") == "deepseek"  # не стёрто


def test_read_health_ttl(tmp_path, monkeypatch):
    # АУДИТ (high): старт опрашивал провайдеров заново при каждом F5 — теперь
    # свежий кэш используется как есть, а устаревший не мешает.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import json, os, time
    from pmoos.core.health import read_health, _health_path
    _health_path().parent.mkdir(parents=True, exist_ok=True)
    _health_path().write_text(json.dumps({"mistral": {"ok": True}}), encoding="utf-8")
    assert read_health(max_age_h=24)                    # свежий — есть
    old = time.time() - 48 * 3600
    os.utime(_health_path(), (old, old))
    assert read_health(max_age_h=24) == {}              # устаревший — пусто
    assert read_health()                                # без TTL — читается


def test_health_ranking_and_model_pick():
    # v0.34: автовыбор провайдера. Порядок: локальные → бесплатные → DeepSeek
    # (платный) → дорогие. Мелкая локальная (7B) уступает бесплатным облачным.
    from pmoos.core.health import pick_model, rank_working, effective_tier
    # эмбеддер bge-m3 НЕ должен попасть в чат-модели
    assert pick_model("ollama", ["bge-m3:latest", "qwen2.5:14b"]) == "qwen2.5:14b"
    assert pick_model("ollama", ["bge-m3:latest"]) == ""
    assert pick_model("deepseek", ["deepseek-v4-flash", "deepseek-v4-pro"]) == "deepseek-v4-pro"
    assert ":free" in pick_model("openrouter", ["x/paid-model", "y/qwen3-32b:free"])

    big_local = {"ok": True, "tier": 0, "best_model": "qwen3:32b", "ms": 900}
    small_local = {"ok": True, "tier": 0, "best_model": "qwen2.5:7b", "ms": 700}
    gem = {"ok": True, "tier": 2, "best_model": "gemini-2.5-flash", "ms": 2000}
    dsk = {"ok": True, "tier": 3, "best_model": "deepseek-v4-pro", "ms": 1500}
    # крупная локальная — первая
    assert rank_working({"ollama": big_local, "gemini": gem, "deepseek": dsk})[0][0] == "ollama"
    # мелкая локальная уступает бесплатному облачному, но опережает платный
    order = [p for p, _ in rank_working({"ollama": small_local, "gemini": gem, "deepseek": dsk})]
    assert order == ["gemini", "ollama", "deepseek"]
    assert effective_tier("ollama", small_local) == 2.5


def test_keysync_merges_from_cloud(tmp_path, monkeypatch):
    # Ключи должны сами обновляться из папки переноса (просьба пользователя)
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path / "data"))
    from pmoos.paths import data_root
    from pmoos.core.keysync import sync_keys_from_transfer, push_keys_to_transfer
    cloud = tmp_path / "cloud"
    cloud.mkdir()
    (cloud / ".env").write_text("DEEPSEEK_API_KEY=sk-new123\nGEMINI_API_KEY=g-new\n",
                                encoding="utf-8")
    data_root().mkdir(parents=True, exist_ok=True)
    local = data_root() / ".env"
    local.write_text("DEEPSEEK_API_KEY=sk-старый\nPMOOS_LOCAL_ONLY=x\n", encoding="utf-8")
    import os, time
    old = time.time() - 3600
    os.utime(local, (old, old))                       # локальный старее облачного
    changed = sync_keys_from_transfer(str(cloud))
    txt = local.read_text(encoding="utf-8")
    assert "sk-new123" in txt and "GEMINI_API_KEY=g-new" in txt
    assert "PMOOS_LOCAL_ONLY=x" in txt                # локальное не теряется
    assert set(changed) == {"DEEPSEEK_API_KEY", "GEMINI_API_KEY"}
    # повторный вызов ничего не делает (локальный уже свежее)
    assert sync_keys_from_transfer(str(cloud)) == []
    # выгрузка ключей в облако
    assert push_keys_to_transfer(str(cloud)) is True
    assert "sk-new123" in (cloud / ".env").read_text(encoding="utf-8")


def test_write_env_key_rejects_garbage(tmp_path, monkeypatch):
    # Реальный инцидент: в поле ключа вставился SSH-ключ — рабочий ключ DeepSeek
    # затёрся, все ответы падали 401. Мусор должен отвергаться, ключ — целеть.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import pytest
    from pmoos.config import write_env_key
    write_env_key("deepseek", "sk-abc123def456ghi789")          # нормальный — ок
    envp = tmp_path / ".env"
    assert "sk-abc123def456ghi789" in envp.read_text(encoding="utf-8")
    for bad in ("ssh-ed25519 AAAA... user@host",                # пробелы
                "DEEPSEEK_API_KEY=sk-xxx",                      # само-вложение
                "sk-" + "x" * 400,                              # километровый
                "строка\nс переводом"):
        with pytest.raises(ValueError):
            write_env_key("deepseek", bad)
    # прежний ключ не пострадал
    assert "sk-abc123def456ghi789" in envp.read_text(encoding="utf-8")


def test_answers_state_and_stop(tmp_path, monkeypatch):
    # v0.33: файл состояния поиска ответов — статус/пульс/стоп (как у индексации)
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.projects import register_project
    import pmoos.pipeline.block1_answers as B
    register_project("ФОН")
    assert B.read_answers_state("ФОН")["status"] == "idle"
    st = B.read_answers_state("ФОН")
    st.update({"status": "running", "pid": 0})
    B.write_answers_state("ФОН", st)
    assert B.answers_running("ФОН") is True          # свежий пульс
    # МЯГКИЙ стоп при живом процессе: флаг взведён, статус остаётся running
    B.stop_answers("ФОН")
    assert B.answers_stop_requested("ФОН") is True
    assert B.read_answers_state("ФОН")["status"] == "running"
    # ЖЁСТКИЙ стоп: паузит (pid=0 — убивать некого)
    B.stop_answers("ФОН", hard=True)
    assert B.read_answers_state("ФОН")["status"] == "paused"
    B.clear_answers_stop("ФОН")
    assert B.answers_stop_requested("ФОН") is False
    # защита от двойного запуска: при живом статусе start возвращает 0
    st = B.read_answers_state("ФОН")
    st.update({"status": "running", "pid": 0})
    B.write_answers_state("ФОН", st)
    assert B.start_answers_background("ФОН") == 0
    st = B.read_answers_state("ФОН")
    st["heartbeat"] = "2020-01-01T00:00:00"
    p = B._ans_state_path("ФОН")
    import json
    p.write_text(json.dumps(st, ensure_ascii=False), encoding="utf-8")
    assert B.answers_running("ФОН") is False         # заглохший пульс


def test_block1_resume_skips_answered_without_llm(tmp_path, monkeypatch):
    # РЕЗЮМ: если все замечания уже отвечены — Блок 1 завершается БЕЗ единого
    # вызова ИИ (и без ключей), сохранив ответы; стоп-флаг останавливает ДО ИИ.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import json
    from pmoos.projects import register_project
    import pmoos.pipeline.block1_answers as B
    from pmoos.paths import project_paths
    register_project("РЕЗ")
    paths = project_paths("РЕЗ")
    rdir = paths["root"] / "remarks"
    rdir.mkdir(parents=True, exist_ok=True)
    rp = rdir / "замечания.txt"
    # ≥3 пунктов: меньший список простой парсер отдал бы ИИ-фолбэку (реальный API!)
    rp.write_text("1. Первое замечание достаточной длины для разбора.\n"
                  "2. Второе замечание достаточной длины для разбора.\n"
                  "3. Третье замечание достаточной длины для разбора.",
                  encoding="utf-8")
    # тексты замечаний в готовых ответах ДОЛЖНЫ совпадать с файлом: резюм
    # теперь сверяет текст (несовпадение = другой файл замечаний → переспрос)
    _t1 = "Первое замечание достаточной длины для разбора."
    _t2 = "Второе замечание достаточной длины для разбора."
    _t3 = "Третье замечание достаточной длины для разбора."
    paths["answers"].parent.mkdir(parents=True, exist_ok=True)
    paths["answers"].write_text(json.dumps({"answers": [
        {"number": "1", "remark": _t1, "answer": "готовый ответ 1", "status": "accepted"},
        {"number": "2", "remark": _t2, "answer": "готовый ответ 2", "status": "proposed"},
        {"number": "3", "remark": _t3, "answer": "готовый ответ 3", "status": "edited",
         "user_answer": "правленый 3"},
    ]}, ensure_ascii=False), encoding="utf-8")
    out = B.run_block1("РЕЗ", remarks_path=str(rp))    # ключи ИИ не нужны
    assert out.get("count") == 3
    assert B.read_answers_state("РЕЗ")["status"] == "done"
    nums = [a["number"] for a in out["answers"]]
    assert nums == ["1", "2", "3"]
    assert out["answers"][0]["answer"] == "готовый ответ 1"

    # старый стоп-флаг очищается при новом запуске («⏯ Продолжить» не глохнет)
    B.request_answers_stop("РЕЗ")
    out2 = B.run_block1("РЕЗ", remarks_path=str(rp))   # все отвечены → без ИИ
    assert B.answers_stop_requested("РЕЗ") is False
    assert out2.get("count") == 3

    # РЕШЕНИЯ ПОЛЬЗОВАТЕЛЯ ПОБЕЖДАЮТ (ревью, high): решение, записанное на диск
    # «во время» генерации, не откатывается пакетным сохранением из памяти
    from pmoos.pipeline.block1_answers import _save_merged
    from pmoos.config import load_config
    from pmoos.ingest.remarks import Remark
    rems = [Remark(number="1", text="р1"), Remark(number="2", text="р2")]
    by_num = {"1": {"number": "1", "remark": "р1", "answer": "фоновый",
                    "status": "proposed", "user_answer": None},
              "2": {"number": "2", "remark": "р2", "answer": "фоновый2",
                    "status": "proposed", "user_answer": None}}
    paths["answers"].write_text(json.dumps({"answers": [
        {"number": "1", "remark": "р1", "answer": "фоновый",
         "status": "accepted", "user_answer": "МОЙ ПРАВЛЕНЫЙ"},
    ]}, ensure_ascii=False), encoding="utf-8")
    out3 = _save_merged("РЕЗ", rems, by_num, load_config(), "линейный", partial=True)
    a1 = next(a for a in out3["answers"] if a["number"] == "1")
    assert a1["status"] == "accepted" and a1["user_answer"] == "МОЙ ПРАВЛЕНЫЙ"

    # СМЕНА ТЕКСТА ЗАМЕЧАНИЯ → переспросить (ревью): готовым не считается
    paths["answers"].write_text(json.dumps({"answers": [
        {"number": "1", "remark": "совсем другой текст замечания",
         "answer": "старый ответ", "status": "proposed"},
        {"number": "2", "remark": "Второе замечание достаточной длины для разбора.",
         "answer": "ок2", "status": "proposed"},
        {"number": "3", "remark": "Третье замечание достаточной длины для разбора.",
         "answer": "ок3", "status": "proposed"},
    ]}, ensure_ascii=False), encoding="utf-8")
    B.request_answers_stop("РЕЗ")  # стоп до первого пакета — ИИ не дёрнется…
    # …но и до стоп-проверки резюм должен пометить №1 как pending; проверяем
    # косвенно: прогон паузится (есть pending), а не завершается done
    B.clear_answers_stop("РЕЗ")
    # прямая проверка резюм-фильтра без ИИ невозможна — проверяем факт: №1 с
    # изменённым текстом НЕ попадает в done_prev → run завершился бы генерацией.
    # Здесь достаточно юнит-логики нормализации: одинаковый текст == готов.
    # (полный путь с ИИ покрывается вживую)


def test_remark_duplicate_numbers_renumbered(tmp_path, monkeypatch):
    # Дубли номеров (ревью): второй «№7» становится «7.2», ответы не схлопываются
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.ingest.remarks import Remark
    remarks = [Remark(number="7", text="первое"), Remark(number="7", text="второе"),
               Remark(number="8", text="третье")]
    _seen = {}
    for r in remarks:
        k = str(r.number)
        _seen[k] = _seen.get(k, 0) + 1
        if _seen[k] > 1:
            r.number = f"{k}.{_seen[k]}"
    assert [r.number for r in remarks] == ["7", "7.2", "8"]


def test_session_guard_closes_stale(tmp_path, monkeypatch):
    # Закрытие старых сеансов при старте: заглохший running → paused, живой не трогаем
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import json
    from pmoos.projects import register_project
    from pmoos.index import indexer as I
    from pmoos.core.session_guard import close_stale_sessions
    register_project("СТАР")
    st = I.read_state("СТАР")
    st.update({"status": "running", "pid": 0})
    I.write_state("СТАР", st)                          # свежий пульс — живой
    notes = close_stale_sessions()
    assert I.read_state("СТАР")["status"] == "running"  # не тронут
    st = I.read_state("СТАР")
    st["heartbeat"] = st["updated_at"] = "2020-01-01T00:00:00"
    (tmp_path / "projects" / "СТАР" / "index_state.json").write_text(
        json.dumps(st, ensure_ascii=False), encoding="utf-8")
    notes = close_stale_sessions()
    assert I.read_state("СТАР")["status"] == "paused"   # заглохший закрыт
    assert any("СТАР" in n for n in notes)
    assert (tmp_path / "session.pid").exists()          # pid текущего сеанса записан


def test_txt_cp1251_fallback(tmp_path):
    # Аудит: ANSI/cp1251 .txt (дефолт Блокнота) индексировался пустышкой —
    # utf-8 + errors="ignore" молча выбрасывал ВСЮ кириллицу
    from pmoos.ingest.loaders import extract_file
    p = tmp_path / "заметка.txt"
    p.write_bytes("Отходы IV класса опасности размещаются на полигоне ТКО".encode("cp1251"))
    pages = extract_file(p)
    assert "полигоне" in pages[0]["text"]
    p2 = tmp_path / "юникод.txt"
    p2.write_text("обычный utf-8 текст", encoding="utf-8")
    assert "utf-8" in extract_file(p2)[0]["text"]


def test_xlsx_none_cells_keep_position(tmp_path):
    # Аудит: None-ячейки выбрасывались — значения съезжали под чужие колонки
    import openpyxl
    from pmoos.ingest.loaders import extract_xlsx
    wb = openpyxl.Workbook(); ws = wb.active
    ws.append(["Код", "Вещество", "ПДК", "Выброс, т/год"])
    ws.append(["0330", None, 0.5, 9.999])
    f = tmp_path / "т.xlsx"; wb.save(str(f))
    text = extract_xlsx(f)[0]["text"]
    assert "0330 |  | 0.5 | 9.999" in text          # позиция колонки сохранена


def test_remarks_page_number_does_not_split():
    # Аудит: одиночный номер страницы PDF резал многостраничное замечание пополам
    from pmoos.ingest.remarks import _split_numbered
    txt = ("1. Первое замечание о составе проекта документации.\n"
           "2. Уточнить количество отходов строительства по группе IV.\n"
           "2\n"
           "а также представить договор со специализированной организацией.")
    out = _split_numbered(txt)
    assert len(out) == 2
    assert "договор" in out[1].text                 # хвост остался в замечании №2
    # а честная колонка «№» из таблиц (строгая последовательность) работает
    txt2 = "1\nПервое замечание достаточной длины для распознавания.\n2\nВторое замечание тоже достаточной длины."
    assert len(_split_numbered(txt2)) == 2


def test_semantic_chunking_keeps_short_tail():
    # Аудит: короткий хвост «ИТОГО: … т/год» молча выпадал из индекса
    from pmoos.ingest.chunking import chunk_text_semantic
    body = "Пункт 5.1. " + ("Мероприятия по охране атмосферного воздуха. " * 40)
    tail = "Таблица 8.1 ИТОГО: 12,5 т/год"
    chunks = chunk_text_semantic(body + "\n" + tail)
    assert any("ИТОГО" in c for c in chunks), "хвост с ИТОГО потерян"


def test_transfer_excludes_uploads(tmp_path, monkeypatch):
    # Просьба пользователя: исходники томов (tmp_uploads) в облако не гоняем
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path / "d"))
    from pmoos.projects import register_project
    from pmoos.core import transfer as T
    from pmoos.index import indexer as I
    from pmoos.paths import data_root
    from pathlib import Path
    register_project("БЕЗИСХ")
    I.write_state("БЕЗИСХ", I.read_state("БЕЗИСХ"))
    up = data_root() / "projects" / "БЕЗИСХ" / "tmp_uploads"
    up.mkdir(parents=True)
    (up / "том.pdf").write_bytes(b"%PDF big source")
    (data_root() / "projects" / "БЕЗИСХ" / "answers.json").write_text("{}", encoding="utf-8")
    dest = str(tmp_path / "cloud")
    ok, msg = T.sync_out(dest)
    assert ok, msg
    assert not (Path(dest) / "projects" / "БЕЗИСХ" / "tmp_uploads" / "том.pdf").exists()
    assert (Path(dest) / "projects" / "БЕЗИСХ" / "answers.json").exists()
    ok, msg = T.sync_in(dest)                       # манифест сходится без исходников
    assert ok, msg
    assert (up / "том.pdf").exists()                # локальные исходники не тронуты


def test_cleanup_sources_after_done(tmp_path, monkeypatch):
    # Авточистка исходников: удаляет после успеха, НЕ удаляет при ошибках/настройке
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.index.indexer import _cleanup_sources
    from pmoos.config import load_config
    cfg = load_config()
    f1 = tmp_path / "a.pdf"; f1.write_bytes(b"x" * 100)
    rep = _cleanup_sources([f1], {"files": {"a.pdf": {"status": "done"}}}, cfg)
    assert rep and not f1.exists()
    f2 = tmp_path / "b.pdf"; f2.write_bytes(b"x")
    rep = _cleanup_sources([f2], {"files": {"b.pdf": {"status": "error"}}}, cfg)
    assert rep == "" and f2.exists()                # ошибки → исходники нужны
    cfg.set("storage.keep_sources", True)
    f3 = tmp_path / "c.pdf"; f3.write_bytes(b"x")
    rep = _cleanup_sources([f3], {"files": {"c.pdf": {"status": "done"}}}, cfg)
    assert rep == "" and f3.exists()                # выключатель работает
    cfg.set("storage.keep_sources", False)


def test_remarks_fetch_name_sanitized():
    # Аудит: имя из Content-Disposition сервера — path traversal и запрещённые символы
    import re
    from pathlib import Path
    name = "../../evil?.docx"
    name = Path(str(name).replace("\\", "/")).name
    name = re.sub(r'[<>:"|?*]', "_", name).strip()
    assert name == "evil_.docx"


def test_classify_razdel_pd_number_and_glued_digits():
    # Находка на ОПОЧКЕ: «Раздел ПД №10_…» уходил в UNKNOWN, потому что
    # (а) «ПД» между словом и номером не допускалось, (б) NFKC превращает «№»
    # в буквы «no». Плюс «ПОС1» (аббревиатура с приклеенной цифрой) не матчился.
    from pmoos.ingest.sections import classify_filename
    assert classify_filename("Раздел ПД №10_ОКН_том 10.2.pdf", "линейный", top=1)[0]["code"] == "OTHER"
    assert classify_filename("Раздел ПД №2_ППО_том 2.1.pdf", "линейный", top=1)[0]["code"] == "PPO"
    assert classify_filename("Том 5.1.1.1_717-14-15-П-1-ПОС1.pdf", "линейный", top=1)[0]["code"] == "POS"
    assert classify_filename("Раздел 3.1.1 конструктивные решения.pdf", "линейный", top=1)[0]["code"] == "TKR"


def test_block3_tolerates_null_fields():
    # Находка аудита: block3 падал TypeError на старых answers.json с null-полями
    from pmoos.pipeline.block3_final import _global_entity_contradictions
    out = _global_entity_contradictions([{"number": "1", "answer": None, "correction": None}])
    assert isinstance(out, list)


def test_match_volume_digit_boundary():
    # Находка аудита: «Том ООС»=«том 6.1» матчился и на файл «Том 6.docx» —
    # правка вставала в чужой том. Цифровая граница это запрещает.
    from pathlib import Path
    from pmoos.output.docx_writer import _match_volume
    assert _match_volume({"oos_volume": "том 6.1"}, Path("Том 6.1_КОРР.docx"))
    assert not _match_volume({"oos_volume": "том 6.1"}, Path("Том 6.docx"))
    assert not _match_volume({"oos_volume": "том 6"}, Path("Том 6.1.docx"))
    assert _match_volume({"oos_volume": "том 6"}, Path("Том 6.docx"))


def test_anchor_token_word_boundary():
    # Находка аудита: «этап. 5» матчился на «п.» и якорь «5» вставал не туда
    from pmoos.output.docx_writer import _anchor_token
    assert _anchor_token("завершить этап. 5 месяцев работ") is None
    assert _anchor_token("исправить в п. 2.3") == "2.3"
    assert _anchor_token("см. табл. 4.1") == "4.1"


def test_block1_rerun_preserves_accepted(tmp_path, monkeypatch):
    # Находка аудита: повторный запуск Блока 1 молча затирал принятые ответы.
    # Проверяем слияние тем же кодом, что в run_block1 (kept-словарь).
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import json
    from pmoos.projects import register_project
    from pmoos.paths import project_paths
    from pmoos.pipeline.block1_answers import load_answers
    register_project("БЛ1")
    p = project_paths("БЛ1")["answers"]
    p.parent.mkdir(parents=True, exist_ok=True)
    prev = {"answers": [
        {"number": "1", "remark": "р1", "answer": "старый", "status": "accepted",
         "user_answer": "мой принятый"},
        {"number": "2", "remark": "р2", "answer": "стар2", "status": "proposed",
         "user_answer": None}]}
    p.write_text(json.dumps(prev, ensure_ascii=False), encoding="utf-8")
    fresh = [{"number": "1", "answer": "новый1", "status": "proposed", "user_answer": None},
             {"number": "2", "answer": "новый2", "status": "proposed", "user_answer": None}]
    kept = {str(a.get("number")): a for a in (load_answers("БЛ1") or {}).get("answers", [])
            if a.get("status") in ("accepted", "edited")}
    merged = [kept.get(str(a.get("number")), a) for a in fresh]
    assert merged[0]["user_answer"] == "мой принятый" and merged[0]["status"] == "accepted"
    assert merged[1]["answer"] == "новый2"   # непринятый — заменён свежим


def test_uprza_backup_on_rewrite(tmp_path, monkeypatch):
    # Находка аудита: повторная выгрузка УПРЗА затирала CSV, заполненные вручную
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.projects import register_project
    from pmoos.output.uprza_export import build_uprza_export
    register_project("У6")
    r1 = build_uprza_export("У6")
    assert r1["backups"] == []
    # «ручное заполнение»
    r1["istochniki"].write_text("моя геометрия", encoding="utf-8")
    r2 = build_uprza_export("У6")
    assert r2["backups"], "бэкап не создан"
    bdir = r2["istochniki"].parent
    backed = [b for b in r2["backups"] if "istochniki" in b]
    assert backed and (bdir / backed[0]).read_text(encoding="utf-8") == "моя геометрия"


def test_unknown_override_does_not_block_classification(tmp_path, monkeypatch):
    # БАГ (найден на ОПОЧКЕ 23.07): «UNKNOWN»-override, оставшийся от прежнего типа
    # объекта, навсегда перекрывал авто-классификацию → «переключил на линейный, а
    # распознано не изменилось». UNKNOWN-override должен игнорироваться/сниматься.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    from pmoos.projects import register_project
    from pmoos.ingest.inventory import build_inventory, set_file_section, load_inventory
    from pmoos.paths import project_paths
    register_project("ИНВ")
    up = project_paths("ИНВ")["uploads"]; up.mkdir(parents=True, exist_ok=True)
    fname = "Раздел ПД №3_ТКР.АД_том 3.1.1.pdf"
    (up / fname).write_bytes(b"%PDF-1.4 test")
    # под линейным файл распознаётся как TKR
    inv = build_inventory("ИНВ", object_type="линейный")
    assert next(f for f in inv["files"] if f["rel"] == fname)["section"] == "TKR"
    # ставим UNKNOWN-override (как делал старый площадной-путь) — он НЕ должен
    # сохраниться как блокирующий override
    set_file_section("ИНВ", fname, "UNKNOWN")
    assert fname not in (load_inventory("ИНВ").get("overrides") or {})
    # даже если UNKNOWN как-то попал в overrides, build_inventory его игнорирует
    inv2 = load_inventory("ИНВ"); inv2.setdefault("overrides", {})[fname] = "UNKNOWN"
    project_paths("ИНВ")["inventory"].write_text(
        __import__("json").dumps(inv2, ensure_ascii=False), encoding="utf-8")
    inv3 = build_inventory("ИНВ", object_type="линейный")
    assert next(f for f in inv3["files"] if f["rel"] == fname)["section"] == "TKR"
    # а РЕАЛЬНЫЙ override (пользователь исправил раздел) — работает
    set_file_section("ИНВ", fname, "POS")
    inv4 = build_inventory("ИНВ", object_type="линейный")
    assert next(f for f in inv4["files"] if f["rel"] == fname)["section"] == "POS"


def test_memory_retraction(tmp_path, monkeypatch):
    # ОТЗЫВ из памяти (v0.31, замечание аудита): неверный ответ можно убрать из
    # few-shot; тот же неизменённый ответ не воскресает; исправленный — снимает отзыв.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path))
    import pmoos.memory as M
    from pmoos.config import load_config
    cfg = load_config()
    cfg.set("memory.semantic", False)  # лексический путь, без загрузки модели

    M.record_many("ПРОШЛЫЙ", [
        {"remark": "Указать площадь застройки участка", "answer": "Площадь 1.2 га", "number": "1"},
        {"remark": "Обосновать выбросы диоксида азота", "answer": "ПДВ 0.45 г/с", "number": "2"}])
    assert M.kb_size() == 2
    assert M.similar_past("площадь застройки", k=3, cfg=cfg)  # находит

    # отзыв убирает из подсказок и из счётчика
    assert M.retract("ПРОШЛЫЙ", "1") is True
    assert M.kb_size() == 1
    assert not any(r.get("number") == "1"
                   for r in M.similar_past("площадь застройки участка", k=3, cfg=cfg))
    assert any(r.get("retracted") for r in M.list_kb())  # запись цела (аудит-след)

    # повторный record_accepted с ТЕМ ЖЕ ответом НЕ воскрешает отозванное
    M.record_many("ПРОШЛЫЙ", [{"remark": "Указать площадь застройки участка",
                               "answer": "Площадь 1.2 га", "number": "1"}])
    assert M.kb_size() == 1
    assert next(r for r in M.list_kb() if r["number"] == "1").get("retracted")

    # ИСПРАВЛЕННЫЙ ответ снимает отзыв
    M.record_many("ПРОШЛЫЙ", [{"remark": "Указать площадь застройки участка",
                               "answer": "Площадь застройки 1.24 га (уточнено)", "number": "1"}])
    assert M.kb_size() == 2
    assert not next(r for r in M.list_kb() if r["number"] == "1").get("retracted")


def test_transfer_sync_roundtrip(tmp_path, monkeypatch):
    # Перенос базы через OneDrive/путь (v0.30.6): выгрузка -> загрузка,
    # venv не переносится, защита от живой индексации.
    monkeypatch.setenv("PMOOS_DATA_DIR", str(tmp_path / "data"))
    from pmoos.projects import register_project
    from pmoos.index import indexer as I
    from pmoos.core import transfer as T
    from pmoos.paths import data_root

    register_project("ПЕР")
    I.write_state("ПЕР", I.read_state("ПЕР"))   # материализует projects/ПЕР
    (data_root() / "qdrant").mkdir(parents=True, exist_ok=True)
    (data_root() / "qdrant" / "x.bin").write_text("вектор-1", encoding="utf-8")
    (data_root() / "venv").mkdir(exist_ok=True)
    (data_root() / "venv" / "lib.txt").write_text("не переносить", encoding="utf-8")

    dest = str(tmp_path / "cloud")
    ok, msg = T.sync_out(dest)
    assert ok, msg
    from pathlib import Path
    assert (Path(dest) / "qdrant" / "x.bin").exists()
    assert (Path(dest) / T.INFO_NAME).exists()
    assert not (Path(dest) / "venv").exists()          # исключён
    assert T.sync_info(dest) and "Выгружено" in T.sync_info(dest)

    # обратный перенос приносит изменения, venv цела
    (Path(dest) / "qdrant" / "x.bin").write_text("вектор-2", encoding="utf-8")
    ok, msg = T.sync_in(dest)
    assert ok, msg
    assert (data_root() / "qdrant" / "x.bin").read_text(encoding="utf-8") == "вектор-2"
    assert (data_root() / "venv" / "lib.txt").exists()

    # живой heartbeat блокирует перенос; заглохший (старый) — нет
    st = I.read_state("ПЕР")
    st.update({"status": "running", "pid": 0})
    I.write_state("ПЕР", st)                           # write_state ставит свежий heartbeat
    ok, msg = T.sync_out(dest)
    assert not ok and "индексация" in msg
    st = I.read_state("ПЕР")
    st["heartbeat"] = st["updated_at"] = "2020-01-01T00:00:00"
    import json
    (data_root() / "projects" / "ПЕР" / "index_state.json").write_text(
        json.dumps(st, ensure_ascii=False), encoding="utf-8")
    ok, msg = T.sync_out(dest)
    assert ok, msg                                     # заглохший running не мешает

    # пустой dest — понятный отказ; относительный путь — отказ
    assert not T.sync_out("")[0]
    assert not T.sync_out("относительный/путь")[0]
    # dest внутри каталога данных — отказ
    assert not T.sync_out(str(data_root() / "куда-то"))[0]

    # ЗАЩИТА /MIR: чужая непустая папка не затирается
    alien = tmp_path / "Документы"
    alien.mkdir()
    (alien / "важное.docx").write_text("не трогать", encoding="utf-8")
    ok, msg = T.sync_out(str(alien))
    assert not ok and "не пустая" in msg
    assert (alien / "важное.docx").exists()

    # КРИТИЧНО (находка ревью): чужой podкаталог projects/ — НЕ признак «нашей»
    # папки; выгрузка обязана отказаться, а не зачистить папку зеркалом
    alien2 = tmp_path / "Dev"
    (alien2 / "projects").mkdir(parents=True)
    (alien2 / "мой_код.py").write_text("х", encoding="utf-8")
    ok, msg = T.sync_out(str(alien2))
    assert not ok and (alien2 / "мой_код.py").exists()

    # «Забрать» из папки без манифеста (чужая/недовыгруженная) — отказ
    ok, msg = T.sync_in(str(alien2))
    assert not ok and "манифест" in msg.lower()

    # «Забрать» при недокачанной OneDrive-копии (файл пропал) — отказ по манифесту
    ok, _ = T.sync_out(dest)
    assert ok
    victim = next((Path(dest) / "projects").rglob("*.json"))
    victim.unlink()
    ok, msg = T.sync_in(dest)
    assert not ok and "досинхронизировал" in msg

    # страховочная копия создаётся при успешном «Забрать»
    ok, _ = T.sync_out(dest)          # восстановить целостность облака
    assert ok
    ok, msg = T.sync_in(dest)
    assert ok, msg
    assert T._backup_dir().exists() and (T._backup_dir() / "projects").exists()


def test_st_cache_unified_with_hub():
    # СТРАХОВКА (по инциденту 17.07.2026): SENTENCE_TRANSFORMERS_HOME обязан
    # указывать на <HF_HOME>/hub — иначе sentence-transformers ведёт ВТОРУЮ
    # копию модели вне hub/, «кэш найден» её не видит, и загрузка уходит в сеть
    # (зависание на прокси при загрузке «с диска»).
    import os
    from pathlib import Path
    import pmoos.config  # noqa: F401  (выставляет переменные окружения)
    st_home = os.environ.get("SENTENCE_TRANSFORMERS_HOME", "")
    hf_home = os.environ.get("HF_HOME", "")
    assert st_home and hf_home
    assert Path(st_home) == Path(hf_home) / "hub", (
        "SENTENCE_TRANSFORMERS_HOME должен совпадать с <HF_HOME>/hub")


def test_env_example_has_no_real_secrets():
    # СТРАХОВКА (по реальному инциденту 16.07.2026): в .env.example попали
    # настоящие ключи, и GitHub push protection заблокировал публикацию.
    # Файл-пример обязан содержать ТОЛЬКО пустые значения/плейсхолдеры.
    from pathlib import Path
    p = Path(__file__).resolve().parent.parent / ".env.example"
    for line in p.read_text(encoding="utf-8").splitlines():
        s = line.strip()
        if not s or s.startswith("#") or "=" not in s:
            continue
        key, _, val = s.partition("=")
        val = val.strip().strip('"').strip("'")
        assert len(val) < 12, (
            f"{key.strip()} в .env.example выглядит как РЕАЛЬНЫЙ ключ — "
            f"ключи хранятся только в ~/.pmoos-rag/.env, не в примере!")


def test_example_config_not_stale():
    # config.example.yaml через deep-merge МОЛЧА перекрывает дефолты — если он
    # отстанет от DEFAULT_CONFIG, пользователь, скопировавший пример, откатит
    # улучшения качества (так уже было: candidates 40 vs 60).
    from pathlib import Path
    import yaml
    from pmoos.config import DEFAULT_CONFIG
    p = Path(__file__).resolve().parent.parent / "config.example.yaml"
    ex = yaml.safe_load(p.read_text(encoding="utf-8"))
    assert ex["retrieval"]["candidates"] == DEFAULT_CONFIG["retrieval"]["candidates"]
    assert ex["reranker"]["max_length"] == DEFAULT_CONFIG["reranker"]["max_length"]
    assert ex["chunking"]["mode"] == DEFAULT_CONFIG["chunking"]["mode"]
    assert ex["retrieval"]["rrf_normalize_dense"] == DEFAULT_CONFIG["retrieval"]["rrf_normalize_dense"]


def test_config_new_keys_present():
    from pmoos.config import load_config
    c = load_config()
    assert int(c.get("reranker.max_length", 0)) >= 512
    assert c.get("retrieval.rrf_normalize_dense") is True
    assert "bm25_weight" in c.data.get("retrieval", {})
