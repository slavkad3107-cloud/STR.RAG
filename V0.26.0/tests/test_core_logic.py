"""Юнит-тесты ЧИСТОЙ (детерминированной) логики СтройПроекта.

Проверяют самые хрупкие правила, которые тихо ломаются при правке регэкспов и
невидимы до неверного результата: коды ЗВ, классификация замечаний, группировка
версий томов, схема ответа ИИ, взвешивание гибридного поиска, якоря правок М5,
состав разделов-источников ПП-87. Ничего тяжёлого (без моделей/сети/GPU)."""
from __future__ import annotations


# ─────────────────────── ЗВ: коды и алиасы ───────────────────────
def test_pollutant_bare_hydrocarbon_alias_removed():
    from pmoos.entities import find_pollutants
    # «непредельные углеводороды C1-C5» НЕ должны получить код предельных 2754
    codes = [p["code"] for p in find_pollutants("выбросы непредельных углеводородов C1-C5")]
    assert "2754" not in codes


def test_pollutant_specific_forms_resolve():
    from pmoos.entities import find_pollutants, pollutant_code
    assert pollutant_code("толуол")[0] == "0621"
    assert pollutant_code("сероводород")[0] == "0333"
    codes = {p["code"] for p in find_pollutants("растворитель: ксилол, толуол, ацетон")}
    assert {"0616", "0621", "1401"} <= codes


def test_pollutant_predelnye_still_2754():
    from pmoos.entities import find_pollutants
    codes = [p["code"] for p in find_pollutants("углеводороды предельные C12-C19")]
    assert "2754" in codes


# ─────────────────────── схема ответа ИИ ───────────────────────
def test_normalize_answer_confidence_and_sources():
    from pmoos.pipeline.block1_answers import _normalize_answer
    r = _normalize_answer({"answer": "  текст  ", "confidence": "высокая",
                           "used_sources": ["1", 3, "х"]})
    assert r["answer"] == "текст"
    assert r["confidence"] == "high"
    assert r["used_sources"] == [1, 3]


def test_normalize_answer_empty_is_low():
    from pmoos.pipeline.block1_answers import _normalize_answer
    r = _normalize_answer({"answer": "", "confidence": "bogus"})
    assert r["confidence"] == "low"
    r2 = _normalize_answer({})
    assert r2["answer"] == "" and r2["used_sources"] == []


# ─────────────────────── классификация замечаний ───────────────────────
def test_classify_remark():
    from pmoos.pipeline.block1_answers import _classify_remark
    assert _classify_remark("Выполнить перерасчёт рассеивания") == "Перерасчёт"
    assert _classify_remark("Привести в соответствие с ГОСТ 17.2.3.02") == "Нормативы"
    assert _classify_remark("Приложить договор на вывоз отходов") == "Доп. документы"


# ─────────────────────── взвешенный RRF ───────────────────────
def test_rrf_weight_lifts_bm25():
    from pmoos.retrieval.hybrid import HybridRetriever
    dense = [{"id": "d1"}]   # только dense, rank0
    bm25 = [{"id": "b1"}]    # только bm25, rank0
    # при равных весах d1 идёт первым (dense-список первый). Большой вес bm25
    # должен вытащить b1 в топ — ровно то, что чинит перекос «BM25 тонет».
    eq = HybridRetriever._rrf([dense, bm25], k=60, weights=[1.0, 1.0])
    assert eq[0]["id"] == "d1"
    fused = HybridRetriever._rrf([dense, bm25], k=60, weights=[1.0, 10.0])
    assert fused[0]["id"] == "b1"


def test_rrf_equal_weight_default():
    from pmoos.retrieval.hybrid import HybridRetriever
    a = [{"id": "x"}, {"id": "y"}]
    b = [{"id": "x"}, {"id": "z"}]
    fused = HybridRetriever._rrf([a, b])  # x встречается в обоих → первый
    assert fused[0]["id"] == "x"


# ─────────────────────── якоря правок М5 ───────────────────────
def test_anchor_token():
    from pmoos.output.docx_writer import _anchor_token
    assert _anchor_token("см. табл. 4.1 раздела") == "4.1"
    assert _anchor_token("уточнить в п. 2.3.5") == "2.3.5"
    assert _anchor_token("общая формулировка без места") is None


def test_iter_all_paragraphs_includes_table_cells():
    # М5: якорь правки часто лежит В ТАБЛИЦЕ; обход должен видеть абзацы ячеек,
    # иначе правки уходят «в конец» вместо места (жалоба пользователя).
    import docx
    from pmoos.output.docx_writer import _iter_all_paragraphs
    d = docx.Document()
    d.add_paragraph("обычный абзац тела")
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "Табл. 4.1 параметры источников"
    texts = [p.text for p in _iter_all_paragraphs(d)]
    assert any("обычный абзац" in x for x in texts)
    assert any("4.1" in x for x in texts)  # абзац из ячейки таблицы найден


def test_match_volume():
    from pathlib import Path
    from pmoos.output.docx_writer import _match_volume
    assert _match_volume({"oos_volume": "том 6.1.docx"}, Path("Том 6.1.docx"))
    # разные тома не должны совпадать по обрезанному stem
    assert not _match_volume({"oos_volume": "том 6.1.docx"}, Path("Том 6.2.docx"))


# ─────────────────────── состав разделов-источников ПП-87 ───────────────────────
def test_ios_emission_sources_included():
    from pmoos.ingest.sections import source_section_codes
    codes = set(source_section_codes("площадной"))
    # ЭОМ/ОВиК/ГС теперь источники выбросов и должны попадать в поиск источников
    assert {"IOS_EOM", "IOS_OV", "IOS_GS"} <= codes


# ─────────────────────── группировка версий томов ───────────────────────
def test_base_name_keeps_volume_number():
    from pmoos.versioning.versions import _base_name
    # «том 6.1» и «том 6.2» — РАЗНЫЕ тома (номер не срезаем)
    assert _base_name("ПМООС том 6.1.docx") != _base_name("ПМООС том 6.2.docx")
    # а «изм.2»/«корр»/дата — версии ОДНОГО документа (срезаются)
    assert _base_name("ПМООС том 6.1 изм.2.docx") == _base_name("ПМООС том 6.1.docx")
    assert _base_name("ПМООС том 6.1 корр 2025-06-01.docx") == _base_name("ПМООС том 6.1.docx")


def test_version_rank_orders():
    from pmoos.versioning.versions import _version_rank
    r0 = _version_rank("ПМООС.docx")[0]
    r_corr = _version_rank("ПМООС корр.docx")[0]
    r_v3 = _version_rank("ПМООС v3.docx")[0]
    assert r0 < r_corr
    assert r_v3 > r0


# ─────────────────────── config merge / get-set ───────────────────────
def test_config_deep_merge_and_get():
    from pmoos.config import Config, _deep_merge
    merged = _deep_merge({"a": {"x": 1, "y": 2}}, {"a": {"y": 9, "z": 3}})
    assert merged == {"a": {"x": 1, "y": 9, "z": 3}}
    c = Config({"retrieval": {"top_k": 8}})
    assert c.get("retrieval.top_k") == 8
    c.set("retrieval.top_k", 12)
    assert c.get("retrieval.top_k") == 12
    assert c.get("nope.missing", "def") == "def"


def test_config_new_keys_present():
    from pmoos.config import load_config
    c = load_config()
    assert int(c.get("reranker.max_length", 0)) >= 512
    assert c.get("retrieval.rrf_normalize_dense") is True
    assert "bm25_weight" in c.data.get("retrieval", {})
