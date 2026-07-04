"""Общие помощники для МОДУЛЯ 5 (формирование .docx/.xlsx).

Содержит:
  • стиль-помощники python-docx (шрифт, заголовки, таблицы с явными ширинами);
  • выбор ответов для вывода (принятые/правленые → иначе предлагаемые);
  • формирование строки «ИСТОЧНИК» (раздел/файл/страница) по требованию М5.
"""
from __future__ import annotations

from typing import Any

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────── выбор ответов ───────────────────────────
def accepted_answers(data: dict) -> tuple[list[dict], bool]:
    """Вернуть (answers, accepted_only).

    Если есть принятые/отредактированные — берём только их (accepted_only=True).
    Иначе — все предложенные, но с флагом accepted_only=False (для предупреждения).
    """
    answers = list(data.get("answers", []))
    acc = [a for a in answers if a.get("status") in ("accepted", "edited")]
    if acc:
        return acc, True
    return answers, False


def final_answer_text(a: dict) -> str:
    """Итоговый текст ответа: правка пользователя приоритетнее предложения ИИ."""
    if a.get("user_answer"):
        return str(a["user_answer"]).strip()
    return str(a.get("answer", "") or "").strip()


def source_ref(a: dict) -> str:
    """Строка ИСТОЧНИК: «раздел … / файл … / стр. …» по первому источнику."""
    srcs = a.get("sources") or []
    if not srcs:
        return "—"
    parts = []
    for s in srcs[:2]:
        sec = s.get("section") or ""
        file = s.get("file") or ""
        loc = s.get("loc") or ""
        piece = []
        if sec:
            piece.append(f"раздел {sec}")
        if file:
            piece.append(f"файл {file}")
        if loc:
            piece.append(str(loc))
        parts.append("; ".join(piece) if piece else "—")
    return " | ".join(parts)


def source_ref_lines(a: dict) -> list[str]:
    """Источник построчно (для ячейки .docx)."""
    srcs = a.get("sources") or []
    out = []
    for s in srcs[:3]:
        sec = s.get("section") or "—"
        file = s.get("file") or "—"
        loc = s.get("loc") or "—"
        out.append(f"раздел: {sec}; файл: {file}; стр./место: {loc}")
    return out or ["—"]


# ─────────────────────────── python-docx стиль ───────────────────────────
def set_default_font(doc, name: str = "Arial", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)
    # кириллица: задаём шрифт и для восточно-европейского набора
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def add_title(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def add_heading(doc, text: str, level: int = 1) -> None:
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x5B)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)


def shade_cell(cell, hex_fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def set_repeat_header(row) -> None:
    tr = row._tr
    trpr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trpr.append(th)


def set_col_widths(table, widths_in: list[float]) -> None:
    """Задать ширины колонок в дюймах для всех строк (надёжно в Word/LibreOffice)."""
    table.autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_in):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def cell_text(cell, text: str, *, bold: bool = False, size: int = 10,
              color: tuple[int, int, int] | None = None) -> None:
    """Записать текст в ячейку, поддерживая переносы строк как отдельные абзацы."""
    cell.text = ""
    p = cell.paragraphs[0]
    first = True
    for line in str(text).split("\n"):
        if not first:
            p = cell.add_paragraph()
        first = False
        r = p.add_run(line)
        r.bold = bold
        r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
