"""Конверсия исходных томов в .docx для внесения правок (v0.46).

Пользователь: «добавить pdf и другие форматы» — тома ООС приходят в PDF,
старом .doc, .rtf/.odt. Правки вносятся только в .docx (python-docx), поэтому
исходник приводится к .docx одним из способов (по убыванию качества):

1) .docx — как есть;
2) LibreOffice (soffice --headless --convert-to docx) — .doc/.rtf/.odt/.pdf;
   PDF импортируется фильтром writer_pdf_import (текст постранично, таблицы
   и оформление частично); ни одной новой python-зависимости;
3) Word ≥ 2007 (COM через PowerShell) — если LibreOffice нет: открывает .doc/
   .rtf/.odt и PDF (Word 2013+), сохраняет .docx с оформлением. Word 2003
   (как на ПК пользователя) PDF не откроет и .docx сохраняет ненадёжно —
   для него этот путь не используется;
4) PDF без обоих — pymupdf: текст постранично в простой .docx (оформление
   теряется — об этом честно сообщаем).

Битую кодировку PDF-шрифтов («Ɂɚɤɚɡɱɢɤ») дальше разруливает decode_garbled.
"""
from __future__ import annotations

import os
import shutil
import subprocess
from pathlib import Path

WORD_FORMATS = {".doc", ".docx", ".rtf", ".odt", ".pdf", ".docm", ".dot", ".dotx"}
_WD_FORMAT_DOCX = 16          # wdFormatXMLDocument (Word 2007+)
_NOWIN = getattr(subprocess, "CREATE_NO_WINDOW", 0)


def soffice_path() -> str | None:
    """LibreOffice, если установлен (PATH или стандартные папки)."""
    p = shutil.which("soffice") or shutil.which("soffice.exe")
    if p:
        return p
    for base in (os.environ.get("ProgramFiles", r"C:\Program Files"),
                 os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)")):
        c = Path(base) / "LibreOffice" / "program" / "soffice.exe"
        if c.exists():
            return str(c)
    return None


def word_version() -> float:
    """Версия Word через COM (0 = нет). 11 = Word 2003, 12 = 2007, 15 = 2013…"""
    try:
        r = subprocess.run(
            ["powershell", "-NoProfile", "-Command",
             "try{$w=New-Object -ComObject Word.Application;$v=$w.Version;$w.Quit();$v}catch{'0'}"],
            capture_output=True, text=True, timeout=40, creationflags=_NOWIN)
        return float((r.stdout or "0").strip().split()[0] or 0)
    except Exception:  # noqa: BLE001
        return 0.0


def _convert_via_soffice(src: Path, dst_dir: Path, timeout: int = 900) -> Path:
    exe = soffice_path()
    if not exe:
        raise RuntimeError("LibreOffice не найден")
    args = [exe, "--headless", "--norestore"]
    if src.suffix.lower() == ".pdf":
        args.append("--infilter=writer_pdf_import")
    args += ["--convert-to", "docx:MS Word 2007 XML", "--outdir", str(dst_dir), str(src)]
    r = subprocess.run(args, capture_output=True, text=True, timeout=timeout,
                       creationflags=_NOWIN)
    out = dst_dir / (src.stem + ".docx")
    if not out.exists():
        raise RuntimeError((r.stderr or r.stdout or "LibreOffice не создал .docx")[:300])
    return out


def _convert_via_word(src: Path, dst: Path, timeout: int = 900) -> None:
    """Word 2007+ COM: открыть → сохранить как .docx (без диалогов)."""
    s, d = str(src).replace("'", "''"), str(dst).replace("'", "''")
    ps = (
        "$ErrorActionPreference='Stop'; $w=New-Object -ComObject Word.Application; "
        "$w.Visible=$false; $w.DisplayAlerts=0; try { "
        f"$doc=$w.Documents.Open('{s}', $false, $true, $false); "
        f"$doc.SaveAs2('{d}', {_WD_FORMAT_DOCX}); $doc.Close($false); 'OK' }} "
        "finally { $w.Quit() }"
    )
    r = subprocess.run(["powershell", "-NoProfile", "-Command", ps],
                       capture_output=True, text=True, timeout=timeout, creationflags=_NOWIN)
    if "OK" not in (r.stdout or "") or not dst.exists():
        raise RuntimeError((r.stderr or r.stdout or "Word не вернул результат")[:300])


def _convert_pdf_fallback(src: Path, dst: Path) -> None:
    """PDF → простой .docx через pymupdf: текст постранично, заголовок «стр. N»."""
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    pdf = fitz.open(str(src))
    try:
        for i, page in enumerate(pdf, start=1):
            h = doc.add_paragraph()
            rh = h.add_run(f"стр. {i}")
            rh.bold = True
            rh.font.size = Pt(9)
            for line in (page.get_text("text") or "").splitlines():
                if line.strip():
                    doc.add_paragraph(line.rstrip())
    finally:
        pdf.close()
    doc.save(str(dst))


def _convert_ocr_text(src: Path, dst: Path) -> int:
    """Скан-PDF без текстового слоя: штатный загрузчик (OCR tesseract, как при
    индексации) → простой .docx постранично. Возвращает число символов."""
    from docx import Document
    from docx.shared import Pt
    from .loaders import extract_file
    pages = extract_file(src, ocr=True)
    doc = Document()
    total = 0
    for pg in pages:
        if pg.get("is_table"):
            continue
        h = doc.add_paragraph()
        rh = h.add_run(str(pg.get("loc", "")))
        rh.bold = True
        rh.font.size = Pt(9)
        for line in (pg.get("text") or "").splitlines():
            if line.strip():
                doc.add_paragraph(line.rstrip())
                total += len(line)
    doc.save(str(dst))
    return total


def _docx_text_len(path: Path) -> int:
    from docx import Document
    d = Document(str(path))
    return sum(len(p.text) for p in d.paragraphs)


# порядок способов по умолчанию: PDF — быстрый текст (секунды) + OCR для
# сканов; LibreOffice для PDF даёт текст в рамках и стоит минуты на большом
# томе (замер: 123 с на 200-КБ файле). .doc/.rtf/.odt — LibreOffice сохраняет
# оформление, затем Word 2007+.
_DEFAULT_PREFER = {".pdf": ("pdf-text", "ocr-text", "soffice")}
_OTHER_PREFER = ("soffice", "word")


def to_docx(src: str | Path, dst_dir: str | Path, *,
            prefer: tuple[str, ...] | None = None) -> dict:
    """Привести исходный том к .docx в dst_dir.

    Возвращает {"path": Path, "method": "copy|soffice|word|pdf-text|ocr-text",
    "note": str}. prefer — порядок способов (в тестах можно принудить)."""
    src, dst_dir = Path(src), Path(dst_dir)
    dst_dir.mkdir(parents=True, exist_ok=True)
    ext = src.suffix.lower()
    if ext not in WORD_FORMATS:
        raise ValueError(f"формат «{ext}» не поддерживается: нужны docx/doc/pdf/rtf/odt")
    dst = dst_dir / (src.stem + ".docx")
    if ext == ".docx":
        if src.resolve() != dst.resolve():
            shutil.copy2(src, dst)
        return {"path": dst, "method": "copy", "note": ""}
    if prefer is None:
        prefer = _DEFAULT_PREFER.get(ext, _OTHER_PREFER)
    errors: list[str] = []
    for method in prefer:
        try:
            if method == "soffice":
                if not soffice_path():
                    raise RuntimeError("LibreOffice не установлен")
                out = _convert_via_soffice(src, dst_dir)
                return {"path": out, "method": "soffice",
                        "note": "сконвертировано LibreOffice"
                                + (" (PDF → текст постранично)" if ext == ".pdf" else "")}
            if method == "word":
                ver = word_version()
                if ver < 12:
                    raise RuntimeError(f"Word {ver:g} не подходит (нужен 2007+)")
                if ext == ".pdf" and ver < 15:
                    raise RuntimeError(f"Word {ver:g} не открывает PDF (нужен 2013+)")
                _convert_via_word(src, dst)
                return {"path": dst, "method": "word",
                        "note": "сконвертировано через Word (с оформлением)"}
            if method == "pdf-text" and ext == ".pdf":
                _convert_pdf_fallback(src, dst)
                if _docx_text_len(dst) < 200:
                    raise RuntimeError("в PDF нет текстового слоя (скан)")
                return {"path": dst, "method": "pdf-text",
                        "note": "PDF переведён в текст без оформления — правки "
                                "встанут по тексту; для точного результата лучше "
                                "Word-том"}
            if method == "ocr-text" and ext == ".pdf":
                n = _convert_ocr_text(src, dst)
                if n < 200:
                    raise RuntimeError("OCR не дал текста")
                return {"path": dst, "method": "ocr-text",
                        "note": "скан-PDF распознан OCR и переведён в текст без "
                                "оформления — проверьте результат"}
        except Exception as e:  # noqa: BLE001
            errors.append(f"{method}: {str(e)[:120]}")
    raise RuntimeError(f"не удалось привести «{src.name}» к .docx — " + "; ".join(errors))
