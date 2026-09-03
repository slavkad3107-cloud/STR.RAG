"""МОДУЛЬ 5 (часть 1): формирование откорректированного раздела ПМООС в .docx.

Архитектурное замечание (важно):
  По требованию пользователя исходные файлы проекта НЕ хранятся приложением
  (см. fix #9 — храним только имя проекта и чанки/токены в RAG-базе). Поэтому
  «откорректированный ПМООС» формируется как профессиональный документ-носитель
  корректировок: для каждого принятого ответа выводится конкретная правка в
  раздел ПМООС со ссылкой на источник (раздел/файл/страница). Если пользователь
  передаёт путь к исходному файлу ПМООС (original_oos_path), его текст
  извлекается и добавляется отдельным приложением для удобства сверки —
  непосредственного слепого переписывания произвольного документа не делаем,
  чтобы не повредить нормоконтроль.

Используется python-docx (работает на машине пользователя без Node.js).
Применяются принципы оформления из docx-skill: шрифт Arial, явные ширины
колонок таблиц (DXA), нумерация средствами Word, без юникод-«буллетов».
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..paths import project_paths
from .common import (
    accepted_answers, final_answer_text, source_ref, set_default_font,
    add_title, add_heading,
)


def _add_intro(doc: Document, project: str, object_type: str, n_corr: int, accepted_only: bool) -> None:
    from .. import VERSION
    p = doc.add_paragraph()
    run = p.add_run(
        f"Документ сформирован автоматически системой {VERSION} "
        f"{datetime.now().strftime('%d.%m.%Y %H:%M')}. "
        f"Проект: «{project}». Тип объекта: {object_type}. "
    )
    run.font.size = Pt(10)
    p2 = doc.add_paragraph()
    note = (
        f"Включено корректировок: {n_corr} (только принятые пользователем)."
        if accepted_only else
        f"Включено корректировок: {n_corr} (ВНИМАНИЕ: показаны предлагаемые ответы — "
        f"ни один пункт ещё не принят пользователем в Модуле 4)."
    )
    r2 = p2.add_run(note)
    r2.font.size = Pt(10)
    r2.italic = True
    if not accepted_only:
        r2.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)


def _add_corrections(doc: Document, answers: list[dict]) -> None:
    add_heading(doc, "1. Корректировки раздела ПМООС по замечаниям экспертизы", level=1)
    if not answers:
        doc.add_paragraph("Принятых корректировок нет.")
        return
    for a in answers:
        num = a.get("number", "?")
        add_heading(doc, f"Замечание №{num}", level=2)

        pr = doc.add_paragraph()
        pr.add_run("Замечание эксперта: ").bold = True
        pr.add_run(a.get("remark", "") or "—")

        corr = (a.get("user_answer") or a.get("correction") or "").strip()
        pc = doc.add_paragraph()
        pc.add_run("Вносимая правка в ПМООС: ").bold = True
        pc.add_run(corr or "—")

        ans = final_answer_text(a)
        if ans:
            pa = doc.add_paragraph()
            pa.add_run("Ответ для экспертизы: ").bold = True
            pa.add_run(ans)

        src = source_ref(a)
        ps = doc.add_paragraph()
        rs = ps.add_run(f"Источник: {src}")
        rs.italic = True
        rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        md = (a.get("missing_data") or "").strip()
        if md:
            pm = doc.add_paragraph()
            rm = pm.add_run(f"Требуется дополнить данными: {md}")
            rm.font.size = Pt(9)
            rm.font.color.rgb = RGBColor(0xB0, 0x60, 0x00)


def _add_original_appendix(doc: Document, original_oos_path: str | Path, cfg) -> None:
    from ..ingest.loaders import extract_file
    try:
        pages = extract_file(Path(original_oos_path), ocr=False)
    except Exception as exc:  # noqa: BLE001
        doc.add_paragraph(f"(Не удалось прочитать исходный ПМООС: {exc})")
        return
    doc.add_page_break()
    add_heading(doc, "Приложение А. Исходный текст раздела ПМООС (для сверки)", level=1)
    note = doc.add_paragraph()
    rn = note.add_run(
        "Ниже приведён извлечённый текст исходного (неоткорректированного) раздела. "
        "Используйте его как основу: примените к нему правки из раздела 1."
    )
    rn.italic = True
    rn.font.size = Pt(9)
    for pg in pages:
        txt = (pg.get("text") or "").strip()
        if not txt:
            continue
        loc = pg.get("loc", "")
        if loc:
            h = doc.add_paragraph()
            rh = h.add_run(str(loc))
            rh.bold = True
            rh.font.size = Pt(9)
            rh.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for para in txt.split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)


def build_corrected_oos_docx(project: str, *, original_oos_path: str | Path | None = None,
                             cfg=None, out_path: str | Path | None = None) -> Path:
    """Сформировать .docx с откорректированным разделом ПМООС.

    Возвращает путь к созданному файлу (по умолчанию в out-папке проекта).
    """
    from ..config import load_config
    cfg = cfg or load_config()
    data = _load_answers(project)
    object_type = data.get("object_type") or cfg.get("object_type", "площадной")

    answers, accepted_only = accepted_answers(data)

    doc = Document()
    set_default_font(doc, "Arial", 11)
    add_title(doc, "ОТКОРРЕКТИРОВАННЫЙ РАЗДЕЛ ПМООС/ООС")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsub = sub.add_run("Перечень мероприятий по охране окружающей среды")
    rsub.font.size = Pt(12)
    rsub.bold = True

    _add_intro(doc, project, object_type, len(answers), accepted_only)
    _add_corrections(doc, answers)

    if original_oos_path:
        _add_original_appendix(doc, original_oos_path, cfg)

    out_dir = project_paths(project)["out"]
    out_dir.mkdir(parents=True, exist_ok=True)
    out = Path(out_path) if out_path else out_dir / f"ПМООС_откорректированный_{project}.docx"
    doc.save(str(out))
    return out


def _load_answers(project: str) -> dict[str, Any]:
    from ..pipeline.block1_answers import load_answers
    return load_answers(project)


# ───────────── №10-11..14: правки ПРЯМО в исходных томах (жёлтым) ─────────────

def _anchor_token(text: str) -> str | None:
    """Маркер места из текста правки: «табл. 4.1», «т. 8.3», «раздел 5», «п. 2.3».

    Падежи учтены («в разделУ/пунктЕ/таблицАХ…») — иначе правка молча уходила
    «в конец» при обычной канцелярской формулировке замечания."""
    import re as _re
    # (?<![а-яёa-z]) — граница слова СЛЕВА: без неё «этап. 5» матчился на «п.»,
    # а «результат. 6» — на «т.» (находка аудита: якорь вставал в чужое место)
    m = _re.search(r"(?<![а-яёa-z])(?:табл(?:иц[аыеуах]{1,2}|\.)?|т\.|разд(?:ел[аеуы]?|\.)?|"
                   r"п(?:ункт[аеуы]?|\.)\.?)"
                   r"\s*№?\s*([\d][\d.]*)", (text or "").lower())
    return m.group(1).rstrip(".") if m else None


def _anchor_re(tok: str):
    """Регэксп поиска якоря с границами по цифрам: «4.1» НЕ должен находиться
    внутри «14.1», «4.12» или даты «04.11.2025» (иначе правка вставала не туда)."""
    import re as _re
    return _re.compile(r"(?<![\d.])" + _re.escape(tok) + r"(?![\d])")


def _find_anchor_paragraph(ptexts, tok: str | None):
    """Первый абзац, содержащий якорный номер (общая логика preview и записи).

    ptexts: список (paragraph|None, lower_text). Возвращает paragraph или None."""
    if not tok:
        return None
    rx = _anchor_re(tok)
    for p, lt in ptexts:
        if rx.search(lt):
            return p
    return None


def _iter_all_paragraphs(container, _seen=None):
    """Все абзацы документа: тело + ячейки таблиц (рекурсивно, включая вложенные).

    python-docx `document.paragraphs` НЕ включает абзацы внутри таблиц. В реальных
    томах ООС нумерованные таблицы/пункты («табл. 4.1», «п. 5.2») почти всегда
    лежат в таблицах — без этого обхода якорь не находится и правки уходят «в конец».
    Объединённые (merged) ячейки дедуплицируются по XML-элементу — иначе один и
    тот же абзац отдавался несколько раз."""
    if _seen is None:
        _seen = set()
    for p in container.paragraphs:
        yield p
    for tbl in getattr(container, "tables", []):
        for row in tbl.rows:
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in _seen:
                    continue
                _seen.add(tc_id)
                yield from _iter_all_paragraphs(cell, _seen)


def _insert_paragraph_after(par, runs):
    """Вставляет новый абзац СРАЗУ ПОСЛЕ par. runs = [(text, bold, yellow)]."""
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn
    from docx.enum.text import WD_COLOR_INDEX
    new_p = par._p.makeelement(qn("w:p"), {})
    par._p.addnext(new_p)
    np = Paragraph(new_p, par._parent)
    for t, bold, hl in runs:
        r = np.add_run(t)
        r.bold = bold
        if hl:
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return np


def _sub_bounded(needle: str, hay: str) -> bool:
    """Подстрока с ЦИФРОВОЙ границей справа: «том 6» НЕ находится в «том 6.1»
    (иначе правка для тома 6.1 вставала и в том 6 — находка аудита; тома у
    пользователя реально нумеруются 6/6.1/6.2)."""
    import re as _re
    # граница: не цифра и не «.цифра» («том 6» ≠ «том 6.1»), но точка
    # РАСШИРЕНИЯ допустима («том 6.3» ∈ «том 6.3.docx») — иначе (05.08) ни один
    # ответ не матчился с томом и все 75 правок уезжали в первый том
    return bool(needle) and _re.search(_re.escape(needle) + r"(?!\.?\d)", hay) is not None


def _volume_tokens(text: str) -> set[str]:
    """Номера томов, названные в тексте: «Том 6.1», «том 6.2,», «тома 6.1–6.3»
    → {'6.1','6.2'}. Диапазон «6.1–6.3» раскрывается по последней цифре."""
    import re as _re
    out: set[str] = set()
    low = (text or "").lower().replace("ё", "е")
    for m in _re.finditer(r"том[аеу]?\s*№?\s*(\d+(?:\.\d+)+)(?:\s*[-–—]\s*(\d+(?:\.\d+)+))?", low):
        a, b = m.group(1), m.group(2)
        out.add(a)
        if b:
            pa, pb = a.rsplit(".", 1), b.rsplit(".", 1)
            if pa[0] == pb[0] and pa[1].isdigit() and pb[1].isdigit():
                for k in range(int(pa[1]), int(pb[1]) + 1):
                    out.add(f"{pa[0]}.{k}")
            else:
                out.add(b)
    return out


def _src_volume_token(src: Path) -> str:
    """Номер тома из имени файла: «Раздел ПД №6_ООС_том 6.1.docx» → '6.1'."""
    import re as _re
    m = _re.search(r"том[аеу]?\s*№?\s*(\d+(?:\.\d+)+)", src.stem.lower())
    if m:
        return m.group(1)
    m = _re.search(r"(\d+\.\d+(?:\.\d+)*)", src.stem)
    return m.group(1) if m else ""


def _match_volume(a: dict, src: Path) -> bool:
    """Относится ли принятый ответ к данному тому.

    Смотрим НЕ ТОЛЬКО служебное поле «Том ООС» (оно часто пусто/неверно), но и
    «где править» и текст замечания: «Том 6.2, п. 4.2.3» → том 6.2. Реальный
    случай (05.09): полтора десятка правок «Том 6.2/6.3» ложились в том 6.1."""
    tok = _src_volume_token(src)
    # ЯВНО названный том («Том 6.2, п. 4.2.3» в «где править» или в замечании)
    # главнее служебного поля «Том ООС»: то поле заполняет поиск по базе и
    # часто указывает не тот том (05.09: правки 6.2/6.3 лежали в 6.1)
    named = _volume_tokens(a.get("edit_location") or "") or _volume_tokens(a.get("remark") or "")
    if named and tok:
        return tok in named
    v = (a.get("oos_volume") or "").lower().strip()
    n, stem = src.name.lower(), src.stem.lower()
    if not v:
        return False
    if _sub_bounded(v, n) or _sub_bounded(n, v) or _sub_bounded(stem, v):
        return True
    vp = Path(v)
    # stem берём ТОЛЬКО если v — имя файла с настоящим расширением
    # (иначе Path("том 6.1").stem == "том 6" и правка утекает в чужой том)
    if vp.suffix.lower() in (".docx", ".doc", ".pdf"):
        return _sub_bounded(vp.stem.lower(), n)
    return False


# ─────────────── v0.45: настоящая корректировка тома ───────────────
# Жалоба пользователя (05.08): «текст замечания просто напечатан поверх ООС —
# нужно в ООС находить, где исправлять, что на что, и делать откорректированный
# том». Диагностика на реальных томах ОПОЧКИ показала ТРИ причины:
#  1) исходные .docx — конверсия из PDF с ИСПОРЧЕННОЙ кодировкой шрифта: в XML
#     «Ɂɚɤɚɡɱɢɤ» вместо «Заказчик» (единое смещение +0x1D6 + 3 спецсимвола);
#     поиск по нормальному тексту не находил НИЧЕГО;
#  2) 33 % текста лежит в текстовых рамках (w:txbxContent) — python-docx их
#     не обходит;
#  3) каждая строка PDF — отдельный абзац (медиана 20 символов, слова разорваны
#     переносами «сель скохозяйственного») — сравнение с одним абзацем бессмысленно.
# Решение: декодер кодировки, обход всех w:p (включая рамки), поиск места по
# ОКНУ соседних строк через символьные n-граммы без пробелов (переносам всё
# равно), замена группы строк на «стало» СТАНДАРТНЫМ шрифтом (в кастомном
# шрифте тома обычная кириллица показалась бы кракозябрами), markdown-таблица
# → настоящая таблица docx, не найденное — компактно в конец без текста ответа.
# Один план для предпросмотра и записи.

_GARBLE_SHIFT = 0x1D6
_GARBLE_EXTRA = {"ʋ": "№", "ɺ": "ё", "ʌ": "/", "Ɫ": "Л"}


def decode_garbled(text: str) -> str:
    """Восстановить кириллицу из «ɡɚɤɚɡɱɢɤ»-кодировки PDF→DOCX конверсий."""
    out = []
    for ch in text or "":
        o = ord(ch)
        if ch in _GARBLE_EXTRA:
            out.append(_GARBLE_EXTRA[ch])
        elif 0x0230 <= o <= 0x02AF:
            d = o + _GARBLE_SHIFT
            out.append(chr(d) if 0x0410 <= d <= 0x044F else ch)
        else:
            out.append(ch)
    return "".join(out)


def garble_ratio(text: str) -> float:
    """Доля «испорченных» символов — признак конверсии из PDF."""
    if not text:
        return 0.0
    bad = sum(1 for ch in text if 0x0230 <= ord(ch) <= 0x02AF)
    return bad / len(text)


def _all_paragraphs(doc):
    """ВСЕ абзацы документа в документном порядке — тело, таблицы, текстовые
    рамки (w:txbxContent): python-docx `paragraphs` рамок не видит, а в
    PDF-конверсиях там треть текста."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    # БЕЗ дедупа по id(el): lxml отдаёт временные прокси, id переиспользуется —
    # такой «seen» молча пропускал абзацы (05.08: метки «терялись», 2 из 15).
    # body.iter и так выдаёт каждый элемент ровно один раз.
    for el in doc.element.body.iter(qn("w:p")):
        yield Paragraph(el, doc)


def _norm(s: str) -> str:
    """Для сравнения: декод, нижний регистр, ё→е, без пробелов и пунктуации —
    переносы строк PDF («свя зи») перестают мешать."""
    import re as _re
    s = decode_garbled(s or "").lower().replace("ё", "е")
    return _re.sub(r"[^а-яa-z0-9]+", "", s)


def _grams(s: str, n: int = 5) -> set[str]:
    return {s[i:i + n] for i in range(len(s) - n + 1)} if len(s) >= n else set()


_STOP_RU = {
    "в", "на", "и", "с", "по", "для", "от", "до", "из", "при", "не", "что", "как",
    "или", "а", "о", "об", "у", "к", "за", "же", "то", "это", "его", "их", "бы",
    "был", "была", "были", "было", "быть", "также", "том", "тома", "раздел",
    "разделе", "пункт", "пункте", "указано", "указан", "указана", "указаны",
    "данные", "данных", "проект", "проекта", "проектом", "объект", "объекта",
    "объекте", "настоящий", "настоящем", "рассматриваемого", "рассматриваемый",
    "согласно", "соответствии", "требований", "требованиями", "приведены",
    "приведен", "приведена", "представлены", "представлен", "представлена",
    "отсутствуют", "отсутствует", "имеются", "имеется", "необходимо", "следует",
    # юридические слова-паразиты: уводили правки в список нормативов
    # («Постановления Правительства Российской Федерации», 05.09)
    "постановление", "постановления", "постановлением", "правительства",
    "правительство", "российской", "федерации", "федеральный", "федерального",
    "закона", "закон", "приказ", "приказа", "приказом", "требования", "статьи",
    "пункта", "пункту", "пунктом", "уточнить", "указать", "представить",
    "привести", "дополнить", "обосновать", "замечание", "замечания", "экспертизы",
}


def _sig_words(text: str) -> set[str]:
    """Значимые слова → префиксы 5 букв (грубый стемминг)."""
    import re as _re
    out: set[str] = set()
    for w in _re.findall(r"[а-яёa-z0-9]+", decode_garbled(text or "").lower()):
        if len(w) < 5 or w in _STOP_RU:
            continue
        out.add(w[:5])
    return out


def _loc_hints(location: str) -> list[str]:
    import re as _re
    hints = _re.findall(r"(?:п(?:ункт[аеуы]?|\.)|табл\w*\.?|разд\w*\.?)\s*№?\s*(\d+(?:\.\d+)*)",
                        (location or "").lower())
    multi = [h for h in hints if "." in h]
    single = [h for h in hints if "." not in h]
    return list(dict.fromkeys(multi + single))


def _is_md_table(text: str) -> bool:
    lines = [l for l in (text or "").splitlines() if l.strip()]
    return len(lines) >= 2 and sum(1 for l in lines if l.count("|") >= 2) >= 2


def _md_table_rows(text: str) -> list[list[str]]:
    import re as _re
    rows = []
    for l in (text or "").splitlines():
        if l.count("|") < 2:
            continue
        cells = [c.strip() for c in l.strip().strip("|").split("|")]
        if all(_re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
            continue
        rows.append(cells)
    return rows


class _Index:
    """Индекс абзацев тома для быстрого нечёткого поиска."""

    def __init__(self, doc):
        self.pars = list(_all_paragraphs(doc))
        self.text = [(p.text or "") for p in self.pars]
        self.norm = [_norm(t) for t in self.text]
        self.words = [_sig_words(t) for t in self.text]
        self.garbled = garble_ratio("".join(self.text[:3000]))
        # КОЛОНТИТУЛЫ: в PDF-конверсии верхний/нижний колонтитул повторяется на
        # каждой странице отдельной строкой («ПРИЛОЖЕНИЕ РАСЧЕТ РАССЕИВАНИЯ…»,
        # шифр тома). Пять правок легли в такую строку (05.09) — исключаем
        # строки, повторяющиеся ≥5 раз, из поиска мест и заголовков
        cnt: dict[str, int] = {}
        for nrm in self.norm:
            if nrm:
                cnt[nrm] = cnt.get(nrm, 0) + 1
        # только ДЛИННЫЕ повторы (≥20 символов без пробелов): короткие ячейки
        # «шт», «м3», числа повторяются сотни раз и не колонтитулы (порог без
        # длины вычёркивал 25 % строк тома — всё уходило «вручную»)
        self.noise: set[int] = {i for i, nrm in enumerate(self.norm)
                                if nrm and 20 <= len(nrm) < 160 and cnt[nrm] >= 5}
        # слово → множество индексов абзацев (для предфильтра кандидатов)
        self.inv: dict[str, set[int]] = {}
        for i, ws in enumerate(self.words):
            if i in self.noise:
                continue
            for w in ws:
                self.inv.setdefault(w, set()).add(i)
        # ЗАГОЛОВКИ РАЗДЕЛОВ ТОМА (для поиска места ПО ТЕМЕ): в PDF-конверсиях
        # заголовки идут ЗАГЛАВНЫМИ и БЕЗ номеров («ОХРАНА НЕДР»), номера
        # пунктов из ответов («п. 3.5.1») в тексте не существуют (05.09).
        # Кандидат: короткая строка, не колонтитул, не конец фразы, после неё
        # идёт текст; ЗАГЛАВНЫМИ — всегда, обычным регистром — если перед ней
        # пустая/короткая строка (подраздел).
        import re as _re
        # heads: (idx_первой_строки, слова, заголовок ЗАГЛАВНЫМИ?, текст)
        self.heads: list[tuple[int, set[str], bool, str]] = []
        n = len(self.text)
        i = 0
        while i < n:
            dt = decode_garbled(self.text[i]).strip()
            L = len(dt)
            ok = (12 <= L <= 120 and i not in self.noise and "\t" not in dt
                  and dt[-1] not in ".,;:" and _re.search(r"[А-ЯЁа-яё]{4,}", dt))
            if not ok:
                i += 1
                continue
            letters = _re.sub(r"[^А-ЯЁа-яёA-Za-z]", "", dt)
            caps = bool(letters) and letters == letters.upper()
            prev_short = i == 0 or len(self.text[i - 1].strip()) < 40
            if not (caps or (dt[0].isupper() and prev_short)):
                i += 1
                continue
            # ЗАГЛАВНЫЙ заголовок часто разорван PDF на 2–3 строки
            # («ВЫБРОСЫ ЗАГРЯЗНЯЮЩИХ» / «ВЕЩЕСТВ В АТМОСФЕРУ») — склеиваем
            j = i + 1
            title = dt
            if caps:
                while j < n:
                    nx = decode_garbled(self.text[j]).strip()
                    nl = _re.sub(r"[^А-ЯЁа-яёA-Za-z]", "", nx)
                    if 3 <= len(nx) <= 120 and nl and nl == nl.upper() and "\t" not in nx \
                            and _re.search(r"[А-ЯЁ]{3,}", nx) and j not in self.noise:
                        title += " " + nx
                        j += 1
                    else:
                        break
            body = sum(1 for q in range(j, min(n, j + 6)) if len(self.text[q]) > 60)
            if body >= 2:
                ws = _sig_words(title)
                if ws:
                    self.heads.append((i, ws, caps, title[:100]))
            i = j
        self._head_pos = [h[0] for h in self.heads]

    def find_topic_heading(self, topic: set[str], *, strict: bool = False
                           ) -> tuple[int, int, float, str]:
        """Заголовок раздела, лучше всего совпадающий с темой ответа (значимые
        слова «где править» + замечания). Возвращает (idx, end_idx, score,
        текст заголовка); idx=-1 если совпадение слабое. Заголовки ЗАГЛАВНЫМИ
        (главы) надёжнее строк обычного регистра — для последних порог выше.
        strict=True (название в кавычках): достаточно и одного точного слова,
        если оно покрывает ≥ половины заголовка."""
        best = (-1, -1, 0.0, "")
        for pos, (i, ws, caps, title) in enumerate(self.heads):
            ov = len(topic & ws)
            if ov == 0:
                continue
            cover = ov / len(ws)
            if strict:
                if ov < 2 and cover < 0.5:
                    continue
            elif ov < 2 or (not caps and (ov < 3 and cover < 0.6)):
                continue
            score = cover + 0.05 * ov + (0.1 if caps else 0.0)
            # слабое совпадение по теме (0.27–0.29 на реальных томах) уводило
            # правку в чужой раздел — лучше честно «вручную»
            if score < 0.35 and not strict:
                continue
            if score > best[2]:
                end = self._head_pos[pos + 1] if pos + 1 < len(self._head_pos) else min(len(self.text), i + 600)
                best = (i, end, score, title[:80])
        return best

    def window_text(self, i: int, k: int) -> str:
        return "".join(self.norm[i:i + k])

    def find(self, was: str, used: set[int], lo: int = 0, hi: int | None = None):
        """Лучшее окно строк [i, i+k) для «было»: (i, k, score)."""
        hi = len(self.pars) if hi is None else hi
        ww = _sig_words(was)
        wn = _norm(was)
        wg = _grams(wn)
        if len(ww) < 2 or len(wg) < 8:
            return -1, 0, 0.0
        # кандидаты: абзацы, где встречаются ≥2 значимых слова «было»
        cnt: dict[int, int] = {}
        for w in ww:
            for i in self.inv.get(w, ()):
                if lo <= i < hi:
                    cnt[i] = cnt.get(i, 0) + 1
        # топ-60 кандидатов по числу совпавших значимых слов: на 64k строк ×
        # 75 ответов полный перебор окон занимал бы минуты
        cands = sorted(cnt, key=lambda i: -cnt[i])[:60]
        best = (-1, 0, 0.0)
        best_dist = 10 ** 9
        target_len = len(wn)
        for c in cands:
            for start in range(max(lo, c - 3), c + 1):
                if start in used or start in self.noise:
                    continue
                acc = ""
                for k in range(1, 12):
                    if start + k > hi or (start + k - 1) in self.noise:
                        break
                    acc += self.norm[start + k - 1]
                    if len(acc) < target_len * 0.5:
                        continue
                    if len(acc) > target_len * 2.2 + 40:
                        break
                    g = _grams(acc)
                    if not g:
                        continue
                    score = len(wg & g) / len(wg)
                    # штраф за окно сильно длиннее «было» (захват чужого текста)
                    if len(acc) > target_len * 1.6:
                        score *= 0.9
                    # при РАВНОМ сходстве берём окно, ближайшее по длине к «было»:
                    # иначе выигрывало окно с лишней строкой сверху (заголовок
                    # «Раздел 10…» затирался заменой — найдено тестом)
                    dist = abs(len(acc) - target_len)
                    if score > best[2] or (score == best[2] and dist < best_dist):
                        best = (start, k, score)
                        best_dist = dist
        return best

    _FRONT_RX = None   # строки состава проекта / оглавления — не заголовки тела

    def _is_front_matter(self, i: int) -> bool:
        """«3.3.2 717/14/15-П-1/ТКР.ЭС» (состав проекта), «3.3.2 …….. 45»
        (оглавление) — такие строки первыми содержат номер пункта и раньше
        перехватывали вставку «по заголовку» (реальный случай 05.09)."""
        import re as _re
        if _Index._FRONT_RX is None:
            _Index._FRONT_RX = _re.compile(
                r"\d+/\d+/\d+-|\.{4,}\s*\d*\s*$|\s\d{1,3}\s*$|^содержание|^оглавление")
        return bool(_Index._FRONT_RX.search(decode_garbled(self.text[i]).strip().lower()))

    def find_heading(self, hint: str):
        """Индекс ЗАГОЛОВКА пункта/таблицы «hint» В ТЕЛЕ тома: короткая строка,
        НАЧИНАЮЩАЯСЯ с номера (не «рис. 3.5.1» посреди текста), не из состава
        проекта/оглавления, и после неё идёт обычный текст (≥2 из следующих 6
        строк длиннее 60 символов). Первое такое вхождение."""
        import re as _re
        # после номера — пробел и слово с ЗАГЛАВНОЙ («3.3.2 Характеристика…»);
        # строки таблиц «1 п п Наименование», списки «5 - Особо большой», текст
        # «7 время в течении…» — не заголовки (реальные промахи 05.09)
        # заголовок бывает и ЗАГЛАВНЫМИ («3.5.1 ИНТЕНСИВНОСТЬ ДВИЖЕНИЯ»)
        rx = _re.compile(r"^(?i:п\.?|пункт|табл\w*\.?|таблица|раздел)?\s*№?\s*"
                         + _re.escape(hint) + r"(?![\d])[.)]?\s+[А-ЯЁA-Z][А-ЯЁа-яёA-Za-z]{2,}")
        n = len(self.text)
        for i, t in enumerate(self.text):
            dt = decode_garbled(t).strip()
            if not dt or len(dt) > 140 or "\t" in dt or not rx.match(dt):
                continue
            if self._is_front_matter(i) or i in self.noise:
                continue
            body = sum(1 for j in range(i + 1, min(n, i + 7)) if len(self.text[j]) > 60)
            if body >= 2:
                return i
        return -1


def plan_corrections(doc, answers: list[dict]) -> tuple[list[dict], "_Index"]:
    """ПЛАН правок для одного тома: где и что менять. Общий для preview и записи.
    Элемент: {number, mode: replace|insert|manual|skip, idx, k, score,
    par_text, shall, location, is_table}."""
    ix = _Index(doc)
    used: set[int] = set()
    plan: list[dict] = []
    for a in answers:
        num = a.get("number", "?")
        # ДЕКОДИРУЕМ поля ответа: ИИ мог скопировать «ɢɧɬɟɧɫɢɜɧɨɫɬɶ» из индекса,
        # собранного до фикса кодировки, — иначе мусор уехал бы в том (05.08)
        shall = decode_garbled((a.get("edit_shall") or a.get("correction") or "").strip())
        was = decode_garbled((a.get("edit_was") or "").strip())
        loc = decode_garbled((a.get("edit_location") or "").strip())
        remark = decode_garbled((a.get("remark") or "").strip())
        e = {"number": num, "mode": "manual", "idx": -1, "k": 0, "score": 0.0,
             "par_text": "", "shall": shall, "location": loc,
             "is_table": _is_md_table(shall), "via": ""}
        if not shall:
            e["mode"] = "skip"
            plan.append(e)
            continue
        # заголовки пунктов/таблиц из «где править» — В ТЕЛЕ тома (не в составе
        # проекта и не в оглавлении); окно поиска — сам пункт (до 400 строк)
        hints = _loc_hints(loc)
        heads = [(h, ix.find_heading(h)) for h in hints]
        heads = [(h, hi_) for h, hi_ in heads if hi_ >= 0]
        i, k, s, via = -1, 0, 0.0, ""
        near = None
        if was:
            # 1) «было» внутри нужного пункта — самое надёжное место
            for h, hi_ in heads:
                i2, k2, s2 = ix.find(was, used, hi_, min(len(ix.pars), hi_ + 400))
                if s2 >= 0.45 and s2 > s:
                    i, k, s, via = i2, k2, s2, f"«было» в п. {h}"
            # 2) по всему тому — только при УВЕРЕННОМ сходстве (замена); при
            #    среднем 0.45–0.55 — вставка после найденного места. Ниже 0.45
            #    в чужой текст не лезем (05.09: 0.35 давало «куда попало»)
            if i < 0:
                i2, k2, s2 = ix.find(was, used)
                if s2 >= 0.55:
                    i, k, s, via = i2, k2, s2, "«было» найдено в томе"
                elif s2 >= 0.50:
                    near = (i2, k2, s2)
        if i >= 0:
            used.update(range(i, i + k))
            e.update(mode="replace", idx=i, k=k, score=round(s, 2), via=via,
                     par_text=decode_garbled(" ".join(ix.text[i:i + k]))[:200])
        elif near:
            i2, k2, s2 = near
            used.update(range(i2, i2 + k2))
            e.update(mode="insert", idx=i2 + k2 - 1, k=1, score=round(s2, 2),
                     via="похоже на «было» (сходство среднее)",
                     par_text=decode_garbled(" ".join(ix.text[i2:i2 + k2]))[:160])
        else:
            # 3) сразу после заголовка нужного пункта/таблицы (в теле тома)
            for h, hi_ in heads:
                if hi_ not in used:
                    e.update(mode="insert", idx=hi_, k=1, via=f"после заголовка п. {h}",
                             par_text=decode_garbled(ix.text[hi_])[:160])
                    used.add(hi_)
                    break
            # 4) ПО ТЕМЕ: раздел тома, чей заголовок совпадает со значимыми
            #    словами «где править» + замечания (в PDF-конверсиях номеров
            #    пунктов нет — только названия). Внутри раздела ещё раз ищем
            #    «было»/цитату замечания (порог ниже: область уже верная);
            #    иначе — сразу после заголовка раздела.
            if e["mode"] == "manual":
                import re as _re2
                # название раздела/таблицы в кавычках из «где править» —
                # самый точный ориентир, пробуем его первым
                quoted = " ".join(_re2.findall(r"[«\"„]([^»\"“]{4,80})[»\"“]", loc))
                ti = -1
                if quoted:
                    ti, tend, ts, ttitle = ix.find_topic_heading(_sig_words(quoted), strict=True)
                if ti < 0:
                    topic = _sig_words(loc) | _sig_words(remark)
                    ti, tend, ts, ttitle = ix.find_topic_heading(topic)
                if ti >= 0:
                    placed = False
                    if was:
                        i4, k4, s4 = ix.find(was, used, ti, tend)
                        if s4 >= 0.40:
                            used.update(range(i4, i4 + k4))
                            e.update(mode="replace" if s4 >= 0.55 else "insert",
                                     idx=i4 if s4 >= 0.55 else i4 + k4 - 1,
                                     k=k4 if s4 >= 0.55 else 1, score=round(s4, 2),
                                     via=f"«было» в разделе «{ttitle[:40]}»",
                                     par_text=decode_garbled(" ".join(ix.text[i4:i4 + k4]))[:160])
                            placed = True
                    if not placed and remark:
                        i5, k5, s5 = ix.find(remark, used, ti, tend)
                        if s5 >= 0.40:
                            used.update(range(i5, i5 + k5))
                            e.update(mode="insert", idx=i5 + k5 - 1, k=1, score=round(s5, 2),
                                     via=f"цитата замечания в разделе «{ttitle[:40]}»",
                                     par_text=decode_garbled(" ".join(ix.text[i5:i5 + k5]))[:160])
                            placed = True
                    if not placed and ti not in used:
                        used.add(ti)
                        e.update(mode="insert", idx=ti, k=1, score=round(ts, 2),
                                 via=f"в раздел «{ttitle[:40]}» (по теме)",
                                 par_text=ttitle)
            # 5) по ТЕКСТУ ЗАМЕЧАНИЯ по всему тому — только при высоком сходстве
            if e["mode"] == "manual" and remark:
                i3, k3, s3 = ix.find(remark, used)
                if s3 >= 0.50:
                    last = i3 + k3 - 1
                    used.update(range(i3, i3 + k3))
                    e.update(mode="insert", idx=last, k=1, score=round(s3, 2),
                             via="по цитате из замечания",
                             par_text=decode_garbled(" ".join(ix.text[i3:i3 + k3]))[:160])
        plan.append(e)
    return plan, ix


def _std_run(run):
    """Стандартный шрифт для ВСТАВЛЯЕМОГО текста: в томах-конверсиях шрифт
    кастомный (глифы по смещённым кодам) — обычная кириллица в нём показалась
    бы кракозябрами."""
    from docx.oxml.ns import qn
    from docx.shared import Pt
    run.font.name = "Times New Roman"
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(k), "Times New Roman")
    run.font.size = Pt(12)


def _apply_plan(doc, plan: list[dict], ix: "_Index") -> dict:
    from docx.enum.text import WD_COLOR_INDEX
    from docx.shared import Pt
    stats = {"replace": 0, "insert": 0, "manual": 0, "skip": 0}
    manual: list[dict] = []

    def _yellow(run):
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def _mark(par, num):
        r = par.add_run(f" [изм. по замечанию №{num}]")
        _std_run(r)
        r.italic = True
        r.font.size = Pt(8)
        _yellow(r)

    def _set_par(par, text):
        for r in list(par.runs):
            r.text = ""
        base = par.runs[0] if par.runs else par.add_run("")
        base.text = text
        _std_run(base)
        _yellow(base)
        return base

    def _table_after(par, rows):
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=ncols)
        try:
            tbl.style = "Table Grid"
        except KeyError:
            # в томах-конверсиях из PDF стандартных стилей нет — рисуем
            # границы вручную (w:tblBorders), иначе таблица без линий
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement
            tpr = tbl._tbl.tblPr
            borders = OxmlElement("w:tblBorders")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(_qn("w:val"), "single")
                b.set(_qn("w:sz"), "4")
                b.set(_qn("w:color"), "000000")
                borders.append(b)
            tpr.append(borders)
        for i, row in enumerate(rows):
            for j in range(ncols):
                cell = tbl.cell(i, j)
                cell.text = row[j] if j < len(row) else ""
                for p in cell.paragraphs:
                    for r in p.runs:
                        _std_run(r)
                        r.font.size = Pt(9)
                        r.bold = (i == 0)
                        _yellow(r)
        par._p.addnext(tbl._tbl)
        # ОБЯЗАТЕЛЬНО абзац ПОСЛЕ таблицы: если таблица оказывается последним
        # элементом ячейки (w:tc), Word считает файл повреждённым и не
        # открывает его (реальный случай: том 6.1, «Таблица 10.2» в ячейке —
        # «откорректированный файл не открыть, пишет ошибка»); в теле документа
        # абзац ещё и не даёт двум таблицам подряд слипнуться в одну
        from docx.oxml.ns import qn as _qn2
        tbl._tbl.addnext(par._p.makeelement(_qn2("w:p"), {}))

    for e in plan:
        mode = e["mode"]
        if mode == "skip":
            stats["skip"] += 1
            continue
        if mode == "manual":
            manual.append(e)
            stats["manual"] += 1
            continue
        par = ix.pars[e["idx"]]
        shall, num = e["shall"], e["number"]
        head_lines = [l for l in shall.splitlines() if l.strip() and l.count("|") < 2]
        body = " ".join(head_lines) if e["is_table"] else shall
        if mode == "replace":
            _set_par(par, body)
            _mark(par, num)
            # остальные строки окна — очищаем (текст перенесён в первую)
            for j in range(e["idx"] + 1, e["idx"] + e["k"]):
                for r in list(ix.pars[j].runs):
                    r.text = ""
        else:  # insert после заголовка
            np = _insert_paragraph_after(par, [])
            _set_par(np, body)
            _mark(np, num)
            par = np
        if e["is_table"]:
            from docx.oxml.ns import qn as _qn
            in_box = any(anc.tag == _qn("w:txbxContent") for anc in par._p.iterancestors())
            if in_box:
                # внутри текстовой рамки настоящую таблицу Word может не открыть —
                # кладём строки текстом «ячейка | ячейка»
                for row in _md_table_rows(shall):
                    np2 = _insert_paragraph_after(par, [])
                    _set_par(np2, " | ".join(row))
                    par = np2
            else:
                _table_after(par, _md_table_rows(shall))
        stats[mode] += 1

    if manual:
        h = doc.add_paragraph()
        hr = h.add_run("ПРАВКИ ПО ЗАМЕЧАНИЯМ, ТРЕБУЮЩИЕ РУЧНОГО ВНЕСЕНИЯ "
                       "(место в томе автоматически не найдено)")
        _std_run(hr)
        hr.bold = True
        for e in manual:
            p = doc.add_paragraph()
            r1 = p.add_run(f"№{e['number']}. ")
            _std_run(r1)
            r1.bold = True
            if e["location"]:
                r2 = p.add_run(f"Где: {e['location']}. ")
                _std_run(r2)
                r2.italic = True
            r3 = p.add_run(e["shall"])
            _std_run(r3)
            _yellow(r3)
    return stats


def repair_structure(doc) -> int:
    """Страховка перед сохранением: OOXML требует, чтобы ячейка таблицы (w:tc),
    текстовая рамка (w:txbxContent) и содержимое sdt заканчивались АБЗАЦЕМ, а
    тело документа — абзацем перед w:sectPr. Нарушение = «Word обнаружил
    нечитаемое содержимое». Добавляет пустые w:p где нужно; возвращает число
    ремонтов (в норме 0)."""
    from docx.oxml.ns import qn
    body = doc.element.body
    fixes = 0
    for tag in ("w:tc", "w:txbxContent", "w:sdtContent"):
        for el in body.iter(qn(tag)):
            kids = [c for c in el if c.tag != qn("w:tcPr")]
            if kids and kids[-1].tag != qn("w:p"):
                el.append(el.makeelement(qn("w:p"), {}))
                fixes += 1
    blocks = [c for c in body if c.tag != qn("w:sectPr")]
    if blocks and blocks[-1].tag != qn("w:p"):
        blocks[-1].addnext(body.makeelement(qn("w:p"), {}))
        fixes += 1
    return fixes


# кэш планов: предпросмотр и запись идут подряд, а разбор тома-конверсии
# (60k строк) стоит ~13 с на том — не делаем его дважды
_PLAN_CACHE: dict[tuple, tuple] = {}
_PLAN_TTL = 600.0


def _answers_key(answers: list[dict]) -> tuple:
    return tuple((str(a.get("number")), a.get("status"),
                  len(a.get("edit_shall") or ""), len(a.get("correction") or ""))
                 for a in answers)


def _plan_for(src, mine: list[dict]):
    """(doc, plan, ix) для тома — из кэша, если том и ответы не менялись."""
    import time as _t
    from docx import Document
    key = (str(src), src.stat().st_mtime, _answers_key(mine))
    hit = _PLAN_CACHE.get(key)
    if hit and _t.time() - hit[3] < _PLAN_TTL:
        return hit[0], hit[1], hit[2]
    doc = Document(str(src))
    plan, ix = plan_corrections(doc, mine)
    # держим не больше 3 томов (по 30 МБ каждый)
    if len(_PLAN_CACHE) >= 3:
        _PLAN_CACHE.pop(next(iter(_PLAN_CACHE)))
    _PLAN_CACHE[key] = (doc, plan, ix, _t.time())
    return doc, plan, ix


def _volume_answers(answers: list[dict], srcs: list, si: int, src) -> list[dict]:
    """Ответы данного тома: по полю «Том ООС» и по «Том X.Y» в тексте ответа.
    Ответ, называющий несколько томов («Том 6.1, Том 6.2, Том 6.3»), идёт в
    КАЖДЫЙ из них; не отнесённые ни к одному тому — в первый."""
    if len(srcs) <= 1:
        return list(answers)
    matched_ids = {id(a) for s2 in srcs for a in answers if _match_volume(a, s2)}
    mine = [a for a in answers if _match_volume(a, src)]
    if si == 0:
        mine += [a for a in answers if id(a) not in matched_ids]
    return mine


def _placed_text(e: dict) -> str:
    via = f" [{e['via']}]" if e.get("via") else ""
    if e["mode"] == "replace":
        return (f"ЗАМЕНА {e['k']} стр. (сходство {int(e['score'] * 100)}%){via}: "
                f"«{e['par_text'][:120]}…»")
    if e["mode"] == "insert":
        return f"ВСТАВКА после «{e['par_text'][:80]}»{via}"
    if e["mode"] == "skip":
        return "пропуск: у ответа нет текста правки («стало»/«правка»)"
    return ("в конец тома, раздел «требуют ручного внесения»"
            + (f" (указано: {e['location'][:70]})" if e["location"] else ""))


def preview_corrections(project: str, sources: list) -> dict:
    """DRY-RUN: тот же план, что и при записи — что и КУДА встанет, с цитатой
    заменяемого фрагмента (декодированной) и оценкой сходства."""
    from docx import Document
    data = _load_answers(project)
    answers = [a for a in data.get("answers", [])
               if a.get("status") in ("accepted", "edited")]
    srcs = [Path(s) for s in sources if s]
    result = {"volumes": [], "total": 0, "accepted": len(answers),
              "stats": {"replace": 0, "insert": 0, "manual": 0, "skip": 0}}
    for si, src in enumerate(srcs):
        mine = _volume_answers(answers, srcs, si, src)
        vol = {"volume": src.name, "changes": [], "answers": len(mine)}
        try:
            doc, plan, ix = _plan_for(src, mine)
        except Exception as e:  # noqa: BLE001
            vol["error"] = f"том не читается ({e}) — запись правок для него не выполнится"
            result["volumes"].append(vol)
            continue
        if ix.garbled > 0.03:
            vol["warning"] = (
                "том — конверсия из PDF с испорченной кодировкой текста; места "
                "найдены по восстановленному тексту, вставки сделаны стандартным "
                "шрифтом. Для чистого результата лучше исходный Word-том.")
        for e in plan:
            vol["changes"].append({"number": e["number"], "placed": _placed_text(e),
                                   "correction": e["shall"][:300], "mode": e["mode"]})
            result["stats"][e["mode"]] = result["stats"].get(e["mode"], 0) + 1
        result["volumes"].append(vol)
        result["total"] += len(plan)
    return result


def write_corrected_volumes(project: str, sources: list) -> tuple[list[Path], list[str]]:
    """Откорректированные тома: план → применение → *_КОРР.docx."""
    from docx import Document
    data = _load_answers(project)
    answers = [a for a in data.get("answers", [])
               if a.get("status") in ("accepted", "edited")]
    srcs = [Path(s) for s in sources if s]
    out_dir = project_paths(project)["out"]
    out_dir.mkdir(parents=True, exist_ok=True)
    outs: list[Path] = []
    failed: list[str] = []
    for si, src in enumerate(srcs):
        mine = _volume_answers(answers, srcs, si, src)
        try:
            doc, plan, ix = _plan_for(src, mine)
        except PermissionError:
            raise
        except Exception as e:  # noqa: BLE001
            failed.append(f"{src.name}: {e}")
            print(f"[m5] ПРОПУЩЕН {src.name}: {e}", flush=True)
            continue
        stats = _apply_plan(doc, plan, ix)
        fixes = repair_structure(doc)
        if fixes:
            print(f"[m5] {src.name}: структура починена ({fixes}) — ячейка/рамка "
                  f"заканчивалась таблицей, Word такой файл не открыл бы", flush=True)
        out = out_dir / f"{src.stem}_КОРР.docx"
        doc.save(str(out))
        # документ мутирован — из кэша вон, иначе повторная запись легла бы
        # поверх уже внесённых правок
        for k in [k for k in _PLAN_CACHE if k[0] == str(src)]:
            _PLAN_CACHE.pop(k, None)
        outs.append(out)
        print(f"[m5] {src.name}: замен {stats['replace']}, вставок {stats['insert']}, "
              f"вручную {stats['manual']} → {out.name}", flush=True)
    if failed and not outs:
        raise RuntimeError("Ни один том не удалось открыть: " + "; ".join(failed)
                           + ". Откройте файлы в Word и пересохраните как .docx.")
    return outs, failed
