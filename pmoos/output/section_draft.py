"""Создание КАРКАСА раздела из БАЗЫ проекта (ТЗ 27.07, п. «ВЫГРУЗКА»).

«Создание по БАЗЕ проекта с довнесением/редактированием необходимых
дополнительных данных». Генерируется рабочий каркас docx:
- структура глав по составу раздела (ПМООС / ИЭИ / ОВОС);
- таблица показателей ИЗ БАЗЫ (реестр ДАННЫХ) с источником каждого значения;
- результаты УПРЗА, если импортированы;
- принятые правки из ОТВЕТОВ, разложенные по главам;
- места, где данных нет, помечены «◈ ВНЕСТИ:» — довнесение по ТЗ.

Это НЕ готовый том: текст глав пишет инженер. Каркас снимает рутину —
структуру, цифры с провенансом и перечень недостающего.
"""
from __future__ import annotations

import re
from pathlib import Path

from ..paths import project_paths

# главы каркаса по целевым разделам
CHAPTERS: dict[str, list[str]] = {
    "OOS": [
        "Результаты оценки воздействия объекта на окружающую среду",
        "Мероприятия по охране атмосферного воздуха",
        "Мероприятия по оценке и снижению шумового воздействия",
        "Мероприятия по охране и рациональному использованию земельных ресурсов и почвенного покрова",
        "Мероприятия по сбору, использованию, обезвреживанию, транспортировке и размещению отходов",
        "Мероприятия по охране недр",
        "Мероприятия по охране объектов растительного и животного мира и среды их обитания",
        "Мероприятия по охране поверхностных и подземных вод",
        "Программа производственного экологического контроля (мониторинга)",
        "Перечень и расчёт затрат на реализацию природоохранных мероприятий",
    ],
    "IEI": [
        "Изученность экологических условий",
        "Краткая характеристика природных и техногенных условий",
        "Почвенно-растительные условия",
        "Животный мир",
        "Хозяйственное использование территории",
        "Социальная сфера",
        "Объекты историко-культурного наследия",
        "Оценка современного экологического состояния территории",
        "Предварительный прогноз возможных неблагоприятных изменений",
        "Рекомендации и предложения по предотвращению вредных последствий",
        "Предложения к программе экологического мониторинга",
    ],
    "OCENKA": [
        "Общие сведения о планируемой (намечаемой) хозяйственной деятельности",
        "Описание возможных видов воздействия на окружающую среду",
        "Описание окружающей среды, которая может быть затронута",
        "Оценка воздействия на окружающую среду альтернативных вариантов",
        "Меры по предотвращению и/или уменьшению негативного воздействия",
        "Предложения по мероприятиям производственного экологического контроля",
        "Выявленные при проведении оценки неопределённости",
        "Обоснование выбора варианта реализации",
        "Резюме нетехнического характера",
    ],
}

# показатель реестра → номер главы ПМООС (куда его значение нужнее всего)
_IND_CHAPTER_OOS = {
    "emission_year": 1, "emission_sec": 1, "sources_count": 1,
    "noise": 2, "szz": 1, "waste_year": 4, "water_use": 7, "water_drain": 7,
    "area_build": 3, "area_plot": 3, "length_route": 0, "duration": 0, "workers": 0,
}


def build_section_draft(project: str, target: str = "OOS",
                        *, out_path: str | Path | None = None) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from ..data import registry as R
    from ..ingest.sections import target_name
    from ..output.common import add_heading, add_title, set_default_font
    from ..output.uprza_import import load_uprza_results
    from ..pipeline.block1_answers import load_answers

    reg = R.load_registry(project)
    inds = reg.get("indicators") or {}
    uprza = load_uprza_results(project)
    answers = (load_answers(project) or {}).get("answers", [])
    accepted = [a for a in answers if a.get("status") in ("accepted", "edited")]
    chapters = CHAPTERS.get(target) or CHAPTERS["OOS"]

    doc = Document()
    set_default_font(doc)
    add_title(doc, f"{target_name(target)} — КАРКАС раздела (проект: {project})")
    p = doc.add_paragraph(
        "Черновик сформирован из базы проекта STR.RAG: показатели с указанием "
        "источника, результаты УПРЗА, принятые правки по замечаниям. Места, "
        "требующие довнесения данных, помечены «◈ ВНЕСТИ:».")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── сводная таблица показателей из базы ──
    add_heading(doc, "Основные показатели (из базы проекта)", level=1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(("Показатель", "Значение", "Ед.", "Источник")):
        tbl.rows[0].cells[i].text = h
    missing: list[str] = []
    for m in R.INDICATORS:
        rec = inds.get(m["key"], {})
        val = str(rec.get("value", "") or "").strip()
        if not val:
            missing.append(m["label"])
            continue
        prov = rec.get("provenance") or {}
        row = tbl.add_row().cells
        row[0].text = m["label"]
        row[1].text = val
        row[2].text = rec.get("unit", m["unit"])
        row[3].text = f"{prov.get('file', '')} {prov.get('loc', '')}".strip() or "—"
    if missing:
        pm = doc.add_paragraph("◈ ВНЕСТИ: нет значений — " + "; ".join(missing))
        pm.runs[0].font.color.rgb = RGBColor(0xB0, 0x30, 0x00)

    # ── результаты УПРЗА ──
    if target == "OOS" and uprza and uprza.get("rows"):
        add_heading(doc, "Результаты расчёта рассеивания (импорт из УПРЗА)", level=1)
        t2 = doc.add_table(rows=1, cols=3)
        t2.style = "Table Grid"
        for i, h in enumerate(("Код ЗВ", "Вещество", "Макс. концентрация, доли ПДК")):
            t2.rows[0].cells[i].text = h
        for r in uprza["rows"][:30]:
            row = t2.add_row().cells
            row[0].text = r["code"]
            row[1].text = r["name"] or "—"
            row[2].text = f"{r['max_pdk']:g}"
        if uprza.get("exceedances"):
            pe = doc.add_paragraph(
                "◈ ВНЕСТИ: обоснование по превышениям (>1 ПДК): " + ", ".join(
                    f"{r['code']} — {r['max_pdk']:.2f} ПДК"
                    for r in uprza["exceedances"][:10]))
            pe.runs[0].font.color.rgb = RGBColor(0xB0, 0x30, 0x00)

    # ── привязка правок к главам: каждая правка попадает РОВНО в одну главу
    # с максимальным пересечением слов; общие слова («мероприятия», «охране»)
    # не считаются — иначе одна правка дублировалась во все главы (ревью)
    _stop = {"мероприятия", "мероприятий", "охране", "охраны", "оценке",
             "оценки", "использованию", "снижению", "предложения", "описание"}

    def _chapter_for(a: dict) -> int | None:
        loc = ((a.get("edit_location") or "") + " " + (a.get("remark") or "")).lower()
        best_ci, best_score = None, 0
        for ci, title in enumerate(chapters, start=1):
            words = [w for w in re.findall(r"[а-яё]{6,}", title.lower())
                     if w not in _stop]
            score = sum(1 for w in words if w[:6] in loc)
            if score > best_score:
                best_ci, best_score = ci, score
        return best_ci

    by_chapter: dict[int, list[dict]] = {}
    unassigned: list[dict] = []
    for a in accepted:
        if not (a.get("edit_shall") or a.get("correction")):
            continue
        ci = _chapter_for(a)
        (by_chapter.setdefault(ci, []) if ci else unassigned).append(a)

    for ci, title in enumerate(chapters, start=1):
        add_heading(doc, f"{ci}. {title}", level=1)
        # показатели, «прописанные» в главу (для ПМООС)
        if target == "OOS":
            for key, ch in _IND_CHAPTER_OOS.items():
                if ch == ci - 1 and str((inds.get(key) or {}).get("value", "")).strip():
                    rec = inds[key]
                    doc.add_paragraph(
                        f"{R._BY_KEY[key]['label']}: {rec['value']} "
                        f"{rec.get('unit', '')} "
                        f"(источник: {(rec.get('provenance') or {}).get('file', '—')}).")
        for a in by_chapter.get(ci, []):
            pa = doc.add_paragraph(
                f"[по замечанию №{a.get('number')}] "
                f"{a.get('edit_shall') or a.get('correction')}")
            pa.runs[0].font.size = Pt(10)
        doc.add_paragraph("◈ ВНЕСТИ: текст главы (пишет инженер).")

    if unassigned:
        add_heading(doc, "Правки, не отнесённые к главам (разложить вручную)", level=1)
        for a in unassigned:
            pa = doc.add_paragraph(
                f"[№{a.get('number')}] {a.get('edit_shall') or a.get('correction')}")
            pa.runs[0].font.size = Pt(10)

    out = Path(out_path) if out_path else (
        project_paths(project)["out"]
        / f"КАРКАС_{target}_{project}.docx".replace("/", "_"))
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out
