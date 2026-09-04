"""ГЕНЕРАЦИЯ РАЗДЕЛА ИИ по исходным данным ДРУГИХ разделов ПД (v0.48).

ТЗ: «по исходным данным из разделов проектной документации (без ООС/ИЭИ/
ОЦЕНКИ) осуществить генерацию этих разделов в автоматическом режиме».

Схема (честная, без выдумок):
  для каждой главы целевого раздела (структура — по эталонам пользователя,
  см. output/section_draft.CHAPTERS) → поиск по базе проекта фрагментов из
  разделов-источников (ТКР/ПОС/КР/ТХ/изыскания…) + показатели из ДАННЫХ →
  ИИ пишет текст главы ТОЛЬКО по этим фрагментам, каждое утверждение с
  пометкой источника [файл, стр.], а чего нет в данных — помечает
  «◈ ВНЕСТИ: …». Итог — docx с главами, источниками и ведомостью пробелов.

Работает ФОНОВЫМ процессом (как ответы на замечания): состояние в
section_gen_state.json (прогресс/пульс/стоп), журнал в section_gen_log.txt.
"""
from __future__ import annotations

import json
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

from ..paths import project_paths, APP_ROOT

_SYS = (
    "Ты — главный инженер проекта, пишешь ГЛАВУ раздела проектной документации "
    "«{target}» (Постановление Правительства РФ № 87) по исходным данным других "
    "разделов проекта. Пиши официальным техническим языком, связным текстом с "
    "абзацами; таблицы — в markdown (| колонка | колонка |). "
    "СТРОГО: используй ТОЛЬКО факты из предоставленных фрагментов и показателей; "
    "после каждого факта ставь ссылку на источник в квадратных скобках вида "
    "[файл, место]. Ничего не выдумывай: если для главы нужных данных нет — "
    "напиши, что именно требуется, строкой «◈ ВНЕСТИ: …» (что, из какого "
    "раздела/документа). Не повторяй текст замечаний экспертизы, не пиши "
    "вступлений и извинений."
)
_USER = (
    "РАЗДЕЛ: {target}\nГЛАВА {n}: {chapter}\n\n"
    "ПОКАЗАТЕЛИ ПРОЕКТА ИЗ БАЗЫ (значение · источник):\n{indicators}\n\n"
    "ФРАГМЕНТЫ РАЗДЕЛОВ-ИСТОЧНИКОВ:\n{fragments}\n\n"
    "Напиши текст главы «{chapter}» (600–1500 слов, если данных достаточно; "
    "если данных мало — коротко и с «◈ ВНЕСТИ»)."
)


# ───────────────────── состояние фонового процесса ─────────────────────
def _state_path(project: str) -> Path:
    return project_paths(project)["root"] / "section_gen_state.json"


def log_path(project: str) -> Path:
    return project_paths(project)["root"] / "section_gen_log.txt"


def _stop_path(project: str) -> Path:
    return project_paths(project)["root"] / "section_gen_stop.flag"


def read_state(project: str) -> dict:
    p = _state_path(project)
    if not p.exists():
        return {"status": "idle", "total": 0, "done": 0, "message": "", "pid": 0}
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:  # noqa: BLE001
        return {"status": "idle", "total": 0, "done": 0, "message": "", "pid": 0}


def write_state(project: str, st: dict) -> None:
    p = _state_path(project)
    p.parent.mkdir(parents=True, exist_ok=True)
    st["heartbeat"] = datetime.now().isoformat(timespec="seconds")
    tmp = p.with_suffix(f".{os.getpid()}.tmp")
    tmp.write_text(json.dumps(st, ensure_ascii=False, indent=2), encoding="utf-8")
    for _ in range(5):
        try:
            tmp.replace(p)
            return
        except PermissionError:
            time.sleep(0.2)
    tmp.replace(p)


def is_running(project: str) -> bool:
    st = read_state(project)
    if st.get("status") not in ("running", "starting"):
        return False
    try:
        age = (datetime.now() - datetime.fromisoformat(st.get("heartbeat") or "")).total_seconds()
    except (ValueError, TypeError):
        return False
    return age < 90


def stop_generation(project: str) -> bool:
    _stop_path(project).write_text("stop", encoding="utf-8")
    return True


def _progress(project: str, total: int, done: int, message: str, status: str = "running",
              **extra) -> None:
    st = read_state(project)
    st.update({"status": status, "total": total, "done": done, "message": message,
               "pid": os.getpid() if status == "running" else st.get("pid", 0)})
    st.update(extra)
    write_state(project, st)


# ───────────────────── сбор данных для главы ─────────────────────
def _indicators_text(project: str) -> str:
    try:
        from ..data import registry as R
        inds = (R.load_registry(project).get("indicators") or {})
    except Exception:  # noqa: BLE001
        return "(показатели не собраны)"
    lines = []
    for m in R.INDICATORS:
        rec = inds.get(m["key"], {})
        v = str(rec.get("value") or "").strip()
        if not v:
            continue
        prov = rec.get("provenance") or {}
        lines.append(f"- {m['label']}: {v} {rec.get('unit', m['unit'])} · "
                     f"[{prov.get('file', '—')}, {prov.get('loc', '')}]")
    return "\n".join(lines) or "(показатели не собраны — соберите во вкладке ДАННЫЕ)"


def _fragments_text(hits: list[dict], limit_chars: int = 9000) -> str:
    out, total = [], 0
    for h in hits:
        pl = h.get("payload") or {}
        txt = re.sub(r"\s+", " ", (h.get("text") or pl.get("text") or "")).strip()
        if not txt:
            continue
        piece = f"[{pl.get('file', '?')}, {pl.get('loc', '')}] {txt[:900]}"
        if total + len(piece) > limit_chars:
            break
        out.append(piece)
        total += len(piece)
    return "\n\n".join(out) or "(фрагменты не найдены)"


def default_retriever(cfg, project: str, object_type: str, target: str) -> Callable[[str], list[dict]]:
    """Поиск по базе проекта по разделам-источникам целевого раздела.
    Возвращает функцию query → hits и объект для close()."""
    from ..retrieval.hybrid import HybridRetriever
    from ..ingest.sections import source_section_codes
    retr = HybridRetriever(cfg)
    srcs = [c for c in source_section_codes(object_type, target) if c != target]

    def run(query: str) -> list[dict]:
        try:
            return retr.batch_search(project, [query], sections=srcs or None,
                                     top=int(cfg.get("gen.top_k", 12)))[0]
        except Exception as e:  # noqa: BLE001
            print(f"[gen] поиск: {e}", flush=True)
            return []
    run.close = retr.close  # type: ignore[attr-defined]
    return run


# ───────────────────── основной проход ─────────────────────
def run_section_gen(project: str, target: str = "OOS", *, cfg=None,
                    object_type: str | None = None,
                    retrieve: Callable[[str], list[dict]] | None = None,
                    chat: Callable[..., str] | None = None) -> Path:
    from ..config import load_config
    from ..ingest.sections import target_name
    from ..output.section_draft import CHAPTERS
    cfg = cfg or load_config()
    if object_type is None:
        try:
            from ..index.indexer import read_state as _rs
            object_type = _rs(project).get("object_type") or "площадной"
        except Exception:  # noqa: BLE001
            object_type = "площадной"
    if chat is None:
        from ..core.ai_providers import chat as _chat
        chat = _chat
    own_retr = retrieve is None
    if retrieve is None:
        retrieve = default_retriever(cfg, project, object_type, target)
    chapters = CHAPTERS.get(target) or CHAPTERS["OOS"]
    tname = target_name(target)
    indicators = _indicators_text(project)
    total = len(chapters)
    results: list[dict] = []
    stop = _stop_path(project)
    if stop.exists():
        stop.unlink()
    try:
        for n, chapter in enumerate(chapters, start=1):
            if stop.exists():
                _progress(project, total, n - 1, f"⏹ Остановлено на главе {n}.", status="paused")
                break
            _progress(project, total, n - 1, f"Глава {n}/{total}: поиск данных — «{chapter[:60]}»")
            query = f"{chapter}. {tname}"
            hits = retrieve(query)
            _progress(project, total, n - 1, f"Глава {n}/{total}: ИИ пишет — «{chapter[:60]}»")
            msgs = [{"role": "system", "content": _SYS.format(target=tname)},
                    {"role": "user", "content": _USER.format(
                        target=tname, n=n, chapter=chapter, indicators=indicators,
                        fragments=_fragments_text(hits))}]
            try:
                text = chat(cfg, msgs, module="module4", max_tokens=3500) or ""
            except Exception as e:  # noqa: BLE001
                text = f"◈ ВНЕСТИ: главу не удалось сгенерировать (ошибка ИИ: {str(e)[:200]})"
            srcs = []
            for h in hits:
                pl = h.get("payload") or {}
                s = f"{pl.get('file', '')} {pl.get('loc', '')}".strip()
                if s and s not in srcs:
                    srcs.append(s)
            results.append({"n": n, "chapter": chapter, "text": text.strip(),
                            "sources": srcs[:12], "hits": len(hits)})
            _progress(project, total, n, f"Глава {n}/{total} готова.")
    finally:
        if own_retr:
            try:
                retrieve.close()  # type: ignore[attr-defined]
            except Exception:  # noqa: BLE001
                pass
    out = _write_docx(project, target, tname, results, indicators)
    st = read_state(project)
    done_all = len(results) == total
    _progress(project, total, len(results),
              (f"Готово: {len(results)} глав → {out.name}" if done_all
               else f"Остановлено: {len(results)}/{total} глав сохранено → {out.name}"),
              status="done" if done_all else "paused", output=str(out))
    return out


def _write_docx(project: str, target: str, tname: str, results: list[dict],
                indicators: str) -> Path:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from ..output.common import add_heading, add_title, set_default_font
    from ..output.docx_writer import _is_md_table, _md_table_rows
    doc = Document()
    set_default_font(doc)
    add_title(doc, f"{tname} — проект раздела, сформированный ИИ ({project})")
    p = doc.add_paragraph(
        "Текст сформирован системой STR.RAG по данным других разделов проектной "
        "документации и показателям из базы проекта. Каждое утверждение снабжено "
        "ссылкой на источник; места, где данных нет, помечены «◈ ВНЕСТИ». Документ "
        "требует проверки инженером и нормоконтроля.")
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    gaps: list[str] = []
    for r in results:
        add_heading(doc, f"{r['n']}. {r['chapter']}", level=1)
        block: list[str] = []
        for line in (r["text"] or "").splitlines():
            if line.strip():
                block.append(line)
            else:
                _flush(doc, block, _is_md_table, _md_table_rows)
                block = []
        _flush(doc, block, _is_md_table, _md_table_rows)
        for line in (r["text"] or "").splitlines():
            if "◈ ВНЕСТИ" in line:
                gaps.append(f"Глава {r['n']}: {line.strip()[:200]}")
        if r["sources"]:
            ps = doc.add_paragraph("Источники: " + "; ".join(r["sources"]))
            ps.runs[0].font.size = Pt(8)
            ps.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    add_heading(doc, "Показатели проекта, использованные при генерации", level=1)
    for line in indicators.splitlines():
        doc.add_paragraph(line.lstrip("- "))
    if gaps:
        add_heading(doc, "Ведомость недостающих данных (◈ ВНЕСТИ)", level=1)
        for g in gaps:
            pg = doc.add_paragraph(g)
            pg.runs[0].font.color.rgb = RGBColor(0xB0, 0x30, 0x00)
    out = project_paths(project)["out"] / f"ГЕНЕРАЦИЯ_{target}_{project}.docx".replace("/", "_")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _flush(doc, block: list[str], is_table, rows_fn) -> None:
    if not block:
        return
    text = "\n".join(block)
    if is_table(text):
        rows = rows_fn(text)
        if rows:
            ncols = max(len(r) for r in rows)
            tbl = doc.add_table(rows=len(rows), cols=ncols)
            try:
                tbl.style = "Table Grid"
            except KeyError:
                pass
            for i, row in enumerate(rows):
                for j in range(ncols):
                    tbl.cell(i, j).text = row[j] if j < len(row) else ""
            doc.add_paragraph()
            return
    for line in block:
        line = re.sub(r"^\s*(#+\s*|\*\*|__)", "", line).strip()
        doc.add_paragraph(line)


# ───────────────────── фоновый запуск ─────────────────────
def start_background(project: str, target: str = "OOS",
                     object_type: str | None = None) -> int:
    import subprocess
    if is_running(project):
        return 0
    try:
        from ..index.indexer import is_running as _idx
        if _idx(project):
            raise RuntimeError("идёт индексация — генерация после её завершения")
    except ImportError:
        pass
    st = read_state(project)
    st.update({"status": "starting", "pid": 0, "total": 0, "done": 0,
               "message": "Запуск фонового процесса…", "target": target, "output": ""})
    write_state(project, st)
    lp = log_path(project)
    lp.parent.mkdir(parents=True, exist_ok=True)
    logf = open(lp, "ab")
    logf.write(f"\n===== {datetime.now().isoformat(timespec='seconds')} генерация "
               f"{target}: {project} =====\n".encode("utf-8"))
    logf.flush()
    env = dict(os.environ, PYTHONIOENCODING="utf-8", PYTHONUNBUFFERED="1")
    args = [sys.executable, "-m", "pmoos.pipeline.section_gen",
            "--project", project, "--target", target]
    if object_type:
        args += ["--object-type", object_type]
    kwargs: dict[str, Any] = {"env": env, "cwd": str(APP_ROOT), "stdout": logf, "stderr": logf}
    if os.name == "nt":
        kwargs["creationflags"] = 0x00000200 | 0x00000008
    else:
        kwargs["start_new_session"] = True
    try:
        proc = subprocess.Popen(args, **kwargs)
    except Exception as e:  # noqa: BLE001
        logf.write(f"не удалось запустить: {e}\n".encode("utf-8"))
        logf.close()
        st.update({"status": "error", "message": f"Не удалось запустить фоновый процесс: {e}"})
        write_state(project, st)
        return 0
    logf.close()
    return proc.pid


def _heartbeat(project: str) -> None:
    import threading

    def beat() -> None:
        while True:
            time.sleep(5)
            try:
                st = read_state(project)
                if int(st.get("pid") or 0) != os.getpid() or st.get("status") != "running":
                    return
                write_state(project, st)
            except Exception:  # noqa: BLE001
                pass
    threading.Thread(target=beat, daemon=True).start()


def _main() -> None:
    import argparse
    import traceback
    ap = argparse.ArgumentParser(description="Фоновая генерация раздела STR.RAG")
    ap.add_argument("--project", required=True)
    ap.add_argument("--target", default="OOS")
    ap.add_argument("--object-type", default=None)
    a = ap.parse_args()
    print(f"[gen] старт: {a.project} / {a.target}, pid={os.getpid()}", flush=True)
    st = read_state(a.project)
    st.update({"status": "running", "pid": os.getpid(), "message": "Подготовка…", "target": a.target})
    write_state(a.project, st)
    _heartbeat(a.project)
    try:
        run_section_gen(a.project, a.target, object_type=a.object_type)
    except Exception as e:  # noqa: BLE001
        traceback.print_exc()
        st = read_state(a.project)
        st.update({"status": "error", "pid": 0, "message": f"Ошибка: {e}"})
        write_state(a.project, st)


if __name__ == "__main__":
    _main()
